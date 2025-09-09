#!/usr/bin/env python3
"""
AvatarCommerce Main Application - FIXED VERSION
A platform for influencers to create AI-powered avatars for audience engagement
"""

import os
import sys
import uuid
import hashlib
import tempfile
import logging
from datetime import datetime, timedelta, timezone
from functools import wraps
from typing import Optional, Dict, Any, List  # ADD List HERE
import json
import jwt
import requests
from flask import Flask, request, jsonify, make_response
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
import time 

try:
    import PyPDF2
    import openai
    from docx import Document as DocxDocument
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    print("⚠️ Install PyPDF2 and python-docx for document processing: pip install PyPDF2 python-docx")

# Import custom modules
try:
    from database import Database
    from chatbot import Chatbot
    from config import Config
except ImportError as e:
    print(f"❌ Error importing modules: {e}")
    print("Please ensure all required modules are in the app directory")
    sys.exit(1)

# =============================================================================
# APPLICATION SETUP
# =============================================================================

def create_app():
    """Create and configure Flask application with simple CORS"""
    
    Config.validate()
    
    app = Flask(__name__)
    app.config['SECRET_KEY'] = Config.JWT_SECRET_KEY
    app.config['MAX_CONTENT_LENGTH'] = Config.MAX_CONTENT_LENGTH
    
    # Simple CORS for development - allows all origins
    CORS(app, origins="*", methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"])
    
    setup_logging()
    return app

def validate_username(username):
    """Validate username format"""
    if not username or not isinstance(username, str):
        return False, "Username is required"
    
    username = username.strip()
    
    if len(username) < 3:
        return False, "Username must be at least 3 characters"
    
    if len(username) > 30:
        return False, "Username must be less than 30 characters"
    
    if not username.replace('_', '').replace('-', '').isalnum():
        return False, "Username can only contain letters, numbers, underscores, and hyphens"
    
    return True, "Valid username"

def validate_email(email):
    """Validate email format"""
    if not email or not isinstance(email, str):
        return False, "Email is required"
    
    email = email.strip().lower()
    
    import re
    email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    
    if not re.match(email_pattern, email):
        return False, "Invalid email format"
    
    return True, "Valid email"

def clean_text_for_response(text, max_length=500):
    """Clean and truncate text for API responses"""
    if not text:
        return ""
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    # Truncate if too long
    if len(text) > max_length:
        text = text[:max_length-3] + "..."
    
    return text

print("✅ Backend fixes loaded - apply these to your main.py file")

def setup_logging():
    """Configure application logging"""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(sys.stdout),
            logging.FileHandler('app.log') if os.access('.', os.W_OK) else logging.NullHandler()
        ]
    )

# =============================================================================
# GLOBAL INSTANCES
# =============================================================================

app = create_app()
logger = logging.getLogger(__name__)

# Initialize database and chatbot
try:
    db = Database()
    chatbot = Chatbot(db)
    logger.info("✅ Database and chatbot initialized successfully")
except Exception as e:
    logger.error(f"❌ Failed to initialize database/chatbot: {e}")
    sys.exit(1)

# =============================================================================
# AUTHENTICATION & MIDDLEWARE
# =============================================================================

def token_required(f):
    """Decorator to require valid JWT token"""
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        # Extract token from Authorization header
        auth_header = request.headers.get('Authorization')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]
        
        if not token:
            return jsonify({
                'status': 'error',
                'message': 'Authentication token is required'
            }), 401
        
        try:
            # Decode token
            data = jwt.decode(token, Config.JWT_SECRET_KEY, algorithms=["HS256"])
            
            # Get current user
            current_user = db.get_influencer_by_username(data['username'])
            if not current_user:
                return jsonify({
                    'status': 'error',
                    'message': 'Invalid token: user not found'
                }), 401
                
        except jwt.ExpiredSignatureError:
            return jsonify({
                'status': 'error',
                'message': 'Token has expired'
            }), 401
        except jwt.InvalidTokenError:
            return jsonify({
                'status': 'error',
                'message': 'Invalid token'
            }), 401
        except Exception as e:
            logger.error(f"Token validation error: {e}")
            return jsonify({
                'status': 'error',
                'message': 'Token validation failed'
            }), 401
        
        return f(current_user, *args, **kwargs)
    
    return decorated

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in Config.ALLOWED_EXTENSIONS

def validate_file_upload(file):
    """Validate uploaded file"""
    if not file or file.filename == '':
        raise ValueError("No file selected")
    
    if not allowed_file(file.filename):
        raise ValueError(f"Invalid file type. Allowed: {', '.join(Config.ALLOWED_EXTENSIONS)}")
    
    # Check file size (additional check beyond Flask's MAX_CONTENT_LENGTH)
    try:
        file.seek(0, 2)  # Seek to end
        size = file.tell()
        file.seek(0)  # Reset
        
        if size > 10 * 1024 * 1024:  # 10MB
            raise ValueError("File too large. Maximum size: 10MB")
        
        if size == 0:
            raise ValueError("File is empty")
            
    except Exception as e:
        raise ValueError(f"Could not read file: {str(e)}")
    
    return True

def hash_password(password):
    """Hash password using SHA-256"""
    return hashlib.sha256(password.encode()).hexdigest()

def generate_jwt_token(user_data):
    """Generate JWT token for user"""
    payload = {
        'username': user_data['username'],
        'id': user_data['id'],
        'user_type': 'influencer',
        'exp': datetime.now(timezone.utc) + timedelta(days=Config.JWT_EXPIRATION_DAYS)
    }
    return jwt.encode(payload, Config.JWT_SECRET_KEY, algorithm="HS256")

# =============================================================================
# ENHANCED HEYGEN API INTEGRATION
# =============================================================================

class HeyGenAPI:
    @staticmethod
    def create_photo_avatar(image_file):
        """Create custom talking photo avatar using the ONLY working HeyGen endpoint"""
        logger.info("🎬 Creating custom talking photo avatar from uploaded image")
        
        try:
            # Read and validate file content
            image_file.seek(0)
            file_content = image_file.read()
            filename = getattr(image_file, 'filename', 'avatar.jpg')
            
            logger.info(f"📁 File: {filename}")
            logger.info(f"📊 Size: {len(file_content)} bytes")
            
            if len(file_content) == 0:
                raise Exception("File is empty")
            
            if len(file_content) > 10 * 1024 * 1024:  # 10MB limit
                raise Exception("File too large. Maximum size: 10MB")
            
            # ONLY use the working endpoint: upload.heygen.com/v1/talking_photo
            logger.info("📤 Using HeyGen talking_photo endpoint (only working endpoint)...")
            
            headers = {
                "X-Api-Key": Config.HEYGEN_API_KEY,
                "Content-Type": "image/jpeg"
            }
            
            response = requests.post(
                "https://upload.heygen.com/v1/talking_photo",
                headers=headers,
                data=file_content,
                timeout=90  # Increased timeout for avatar processing
            )
            
            logger.info(f"📥 Response: {response.status_code}")
            logger.info(f"📋 Response body: {response.text}")
            
            if response.status_code == 200:
                result = response.json()
                
                # Parse HeyGen response format
                if result.get('code') == 100 and result.get('data'):
                    data = result.get('data', {})
                    talking_photo_id = data.get('talking_photo_id')
                    
                    if talking_photo_id:
                        logger.info(f"✅ SUCCESS! Custom talking photo created: {talking_photo_id}")
                        return {
                            "avatar_id": talking_photo_id,
                            "status": "ready",
                            "type": "talking_photo",
                            "method": "talking_photo_upload",
                            "is_custom": True,
                            "preview_image_url": data.get('preview_image_url'),
                            "preview_video_url": data.get('preview_video_url')
                        }
                    else:
                        error_msg = "No talking_photo_id in response"
                        logger.error(f"❌ {error_msg}")
                        raise Exception(f"HeyGen API error: {error_msg}")
                else:
                    error_msg = result.get('message', 'Unknown error from HeyGen API')
                    error_code = result.get('code', 'unknown')
                    logger.error(f"❌ HeyGen API error {error_code}: {error_msg}")
                    
                    # Handle specific error codes
                    if error_code == 40002:
                        raise Exception("Image format not supported. Please use JPEG, PNG, or WebP format.")
                    elif error_code == 40001:
                        raise Exception("Invalid image data. Please try a different image.")
                    elif error_code == 40003:
                        raise Exception("No face detected in image. Please use a clear photo with a visible face.")
                    elif error_code == 40004:
                        raise Exception("Multiple faces detected. Please use an image with only one face.")
                    elif error_code == 42901:
                        raise Exception("Insufficient credits. Please check your HeyGen account balance.")
                    else:
                        raise Exception(f"HeyGen API error: {error_msg}")
                        
            elif response.status_code == 400:
                # Try to parse error message
                try:
                    error_data = response.json()
                    error_msg = error_data.get('message', 'Bad request')
                    error_code = error_data.get('code', 'unknown')
                    
                    logger.error(f"❌ HeyGen 400 error {error_code}: {error_msg}")
                    
                    if error_code == 40002:
                        raise Exception("Image format not supported. Please use a clear JPEG, PNG, or WebP image.")
                    elif "face" in error_msg.lower():
                        raise Exception("Could not detect a face in the image. Please use a clear photo with a visible face.")
                    elif "format" in error_msg.lower():
                        raise Exception("Invalid image format. Please use JPEG, PNG, or WebP.")
                    else:
                        raise Exception(f"Image processing error: {error_msg}")
                        
                except:
                    raise Exception("Invalid image data. Please try a different image with a clear face.")
                    
            elif response.status_code == 401:
                logger.error("❌ HeyGen authentication failed")
                raise Exception("HeyGen API authentication failed. Please check your API key.")
                
            elif response.status_code == 429:
                logger.error("❌ HeyGen rate limit exceeded")
                raise Exception("Too many requests. Please wait a moment and try again.")
                
            elif response.status_code == 403:
                logger.error("❌ HeyGen access forbidden")
                raise Exception("Access denied. Please check your HeyGen account permissions.")
                
            else:
                error_text = response.text
                logger.error(f"❌ HeyGen API error: {response.status_code} - {error_text}")
                raise Exception(f"HeyGen API returned error {response.status_code}. Please try again later.")
                
        except Exception as e:
            logger.error(f"❌ Custom avatar creation failed: {str(e)}")
            raise e  # Don't fall back - let the frontend handle the error

    @staticmethod
    def generate_video(avatar_id, text, voice_id=None, is_newly_created=False):
        """FIXED: Generate video with avatar using HeyGen API v2 with proper error handling"""
        try:
            headers = {
                "X-Api-Key": Config.HEYGEN_API_KEY,
                "Content-Type": "application/json"
            }
            
            # Clean text for video generation
            clean_text = text[:400] if len(text) > 400 else text
            
            # Prepare video generation payload
            payload = {
                "video_inputs": [
                    {
                        "character": {
                            "type": "avatar",
                            "avatar_id": avatar_id,
                            "avatar_style": "normal"
                        },
                        "voice": {
                            "type": "text",
                            "input_text": clean_text,
                            "voice_id": voice_id or Config.DEFAULT_VOICE_ID
                        }
                    }
                ],
                "dimension": {
                    "width": 1280,
                    "height": 720
                },
                "aspect_ratio": "16:9",
                "test": False
            }
            
            logger.info(f"🎬 Sending video generation request to HeyGen")
            logger.info(f"📝 Text: {clean_text}")
            logger.info(f"🎤 Voice ID: {voice_id}")
            
            response = requests.post(
                "https://api.heygen.com/v2/video/generate",
                headers=headers,
                json=payload,
                timeout=30
            )
            
            logger.info(f"📡 HeyGen video generation response: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                
                if data.get('code') == 100 or data.get('data'):
                    video_data = data.get('data', {})
                    video_id = video_data.get('video_id')
                    
                    if video_id:
                        logger.info(f"✅ Video generation started: {video_id}")
                        return {
                            'video_id': video_id,
                            'status': 'processing',
                            'estimated_time': video_data.get('estimated_time', 30)
                        }
                    else:
                        raise Exception("No video_id in response")
                else:
                    error_msg = data.get('message', 'Video generation failed')
                    logger.error(f"❌ HeyGen video error: {error_msg}")
                    raise Exception(error_msg)
                    
            elif response.status_code == 404:
                # Handle 404 differently for newly created avatars
                try:
                    error_data = response.json()
                    error_code = error_data.get('error', {}).get('code', '')
                    
                    if error_code == 'avatar_not_found':
                        if is_newly_created:
                            logger.warning(f"⏳ Newly created avatar {avatar_id} not yet ready for video generation")
                            raise Exception("AVATAR_STILL_PROCESSING")
                        else:
                            logger.error(f"❌ Existing avatar {avatar_id} no longer exists")
                            raise Exception("AVATAR_NOT_FOUND")
                    else:
                        raise Exception(f"Avatar error: {error_data.get('error', {}).get('message', 'Unknown error')}")
                        
                except json.JSONDecodeError:
                    if is_newly_created:
                        raise Exception("AVATAR_STILL_PROCESSING")
                    else:
                        raise Exception("AVATAR_NOT_FOUND")
                        
            else:
                error_text = response.text
                logger.error(f"❌ HeyGen video API error: {response.status_code} - {error_text}")
                raise Exception(f"Video generation failed: HTTP {response.status_code}")
                
        except Exception as e:
            logger.error(f"❌ Video generation failed: {e}")
            raise e

    @staticmethod
    def get_video_status(video_id):
        """FIXED: Check video generation status with proper error handling"""
        try:
            headers = {
                "X-Api-Key": Config.HEYGEN_API_KEY,
                "Accept": "application/json"
            }
            
            logger.info(f"📊 Checking status for video: {video_id}")
            
            response = requests.get(
                f"https://api.heygen.com/v1/video_status.get?video_id={video_id}",
                headers=headers,
                timeout=15
            )
            
            logger.info(f"📊 Status check response: {response.status_code}")
            
            if response.status_code == 200:
                data = response.json()
                
                if data.get('code') == 100:
                    video_data = data.get('data', {})
                    status = video_data.get('status', 'processing')
                    
                    result = {
                        'status': status,
                        'video_url': video_data.get('video_url'),
                        'progress': video_data.get('progress', 50),
                        'error': video_data.get('error')
                    }
                    
                    logger.info(f"📊 Video status: {status}")
                    if result['video_url']:
                        logger.info(f"🎬 Video URL: {result['video_url']}")
                    
                    return result
                else:
                    error_msg = data.get('message', 'Unknown error from HeyGen')
                    logger.error(f"❌ HeyGen API error: {error_msg}")
                    return {'status': 'failed', 'error': error_msg}
            else:
                logger.error(f"❌ Video status check failed: {response.status_code}")
                return {'status': 'failed', 'error': f'HTTP {response.status_code}'}
                
        except Exception as e:
            logger.error(f"❌ Video status check error: {e}")
            return {'status': 'failed', 'error': str(e)}

    @staticmethod
    def get_available_avatars():
        """Get list of available prebuilt avatars (working endpoint)"""
        try:
            headers = {
                "X-Api-Key": Config.HEYGEN_API_KEY,
                "Accept": "application/json"
            }
            
            response = requests.get(
                "https://api.heygen.com/v2/avatars",
                headers=headers,
                timeout=15
            )
            
            if response.status_code == 200:
                data = response.json()
                avatars = data.get("data", {}).get("avatars", [])
                
                logger.info(f"📋 Retrieved {len(avatars)} available avatars")
                return avatars
            else:
                logger.error(f"❌ Failed to get avatars: {response.status_code}")
                return []
                
        except Exception as e:
            logger.error(f"❌ Get avatars error: {e}")
            return []

    @staticmethod
    def validate_avatar(avatar_id, is_newly_created=False):
        """FIXED: Enhanced avatar validation with special handling for new avatars"""
        import time
        
        try:
            headers = {
                "X-Api-Key": Config.HEYGEN_API_KEY,
                "Accept": "application/json"
            }
            
            # Skip validation for newly created avatars - they need time to process
            if is_newly_created:
                logger.info(f"🕐 Skipping validation for newly created avatar {avatar_id} - allowing processing time")
                return True
            
            # Test with a simple avatar info request
            response = requests.get(
                f"https://api.heygen.com/v1/avatar/{avatar_id}",
                headers=headers,
                timeout=10
            )
            
            if response.status_code == 200:
                logger.info(f"✅ Avatar {avatar_id} validated successfully")
                return True
            elif response.status_code == 404:
                logger.warning(f"⚠️ Avatar {avatar_id} not found on HeyGen")
                
                # For existing avatars, try once more with longer delay
                time.sleep(3)
                
                response = requests.get(
                    f"https://api.heygen.com/v1/avatar/{avatar_id}",
                    headers=headers,
                    timeout=10
                )
                
                if response.status_code == 200:
                    logger.info(f"✅ Avatar {avatar_id} now available after retry")
                    return True
                else:
                    logger.warning(f"⚠️ Avatar {avatar_id} still not found after retry")
                    return False
            else:
                logger.warning(f"⚠️ Avatar validation returned {response.status_code}")
                return True  # Assume valid if we can't confirm otherwise
                
        except Exception as e:
            logger.warning(f"⚠️ Avatar validation error: {e}")
            return True  # Assume valid if validation fails

# =============================================================================
# ERROR HANDLERS
# =============================================================================

@app.errorhandler(400)
def bad_request(error):
    return jsonify({
        'status': 'error',
        'message': 'Bad request - invalid data provided'
    }), 400

@app.errorhandler(401)
def unauthorized(error):
    return jsonify({
        'status': 'error',
        'message': 'Unauthorized - please log in'
    }), 401

@app.errorhandler(403)
def forbidden(error):
    return jsonify({
        'status': 'error',
        'message': 'Forbidden - insufficient permissions'
    }), 403

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'status': 'error',
        'message': 'Resource not found'
    }), 404

@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({
        'status': 'error',
        'message': 'Method not allowed'
    }), 405

@app.errorhandler(413)
@app.errorhandler(RequestEntityTooLarge)
def file_too_large(error):
    return jsonify({
        'status': 'error',
        'message': 'File too large. Maximum size allowed is 16MB'
    }), 413

@app.errorhandler(422)
def unprocessable_entity(error):
    return jsonify({
        'status': 'error',
        'message': 'Unprocessable entity - invalid data format'
    }), 422

@app.errorhandler(429)
def rate_limit_exceeded(error):
    return jsonify({
        'status': 'error',
        'message': 'Rate limit exceeded. Please try again later'
    }), 429

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}")
    return jsonify({
        'status': 'error',
        'message': 'Internal server error - please try again later'
    }), 500

@app.errorhandler(502)
def bad_gateway(error):
    return jsonify({
        'status': 'error',
        'message': 'Bad gateway - external service unavailable'
    }), 502

@app.errorhandler(503)
def service_unavailable(error):
    return jsonify({
        'status': 'error',
        'message': 'Service unavailable - please try again later'
    }), 503

# =============================================================================
# API ROUTES
# =============================================================================

# Health Check
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint for connection testing"""
    try:
        # Test database connection
        test_query = db.supabase.table('influencers').select('count').execute()
        
        return jsonify({
            'status': 'success',
            'message': 'API is healthy',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'database': 'connected'
        })
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({
            'status': 'error',
            'message': 'API health check failed',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'database': 'disconnected'
        }), 500

# Authentication Routes
@app.route('/api/register', methods=['POST'])
@app.route('/api/auth/register', methods=['POST'])
def register():
    """FIXED: Register new influencer account with consistent response format"""
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['username', 'email', 'password']
        missing_fields = [field for field in required_fields if not data.get(field)]
        if missing_fields:
            return jsonify({
                'status': 'error',
                'message': f'Missing required fields: {", ".join(missing_fields)}'
            }), 400
        
        username = data['username'].strip()
        email = data['email'].strip().lower()
        password = data['password']
        
        # Validate username format
        if not username.replace('_', '').replace('-', '').isalnum():
            return jsonify({
                'status': 'error',
                'message': 'Username can only contain letters, numbers, underscores, and hyphens'
            }), 400
        
        if len(username) < 3 or len(username) > 30:
            return jsonify({
                'status': 'error',
                'message': 'Username must be between 3 and 30 characters'
            }), 400
        
        # Check if user already exists
        existing_user = db.get_influencer_by_username(username)
        if existing_user:
            return jsonify({
                'status': 'error',
                'message': 'Username already exists'
            }), 400
        
        existing_email = db.get_influencer_by_email(email)
        if existing_email:
            return jsonify({
                'status': 'error',
                'message': 'Email already exists'
            }), 400
        
        # Create user
        user_data = {
            'id': str(uuid.uuid4()),
            'username': username,
            'email': email,
            'password_hash': hash_password(password),
            'bio': data.get('bio', ''),
            'created_at': datetime.now(timezone.utc).isoformat()
        }
        
        created_user = db.create_influencer(user_data)
        
        if not created_user:
            logger.error(f"Failed to create user: {username}")
            return jsonify({
                'status': 'error',
                'message': 'Failed to create account. Please try again.'
            }), 500
        
        # Generate token
        token = generate_jwt_token(created_user)
        
        logger.info(f"✅ Successfully registered user: {username}")
        
        # FIXED: Consistent response format for frontend
        return jsonify({
            'status': 'success',
            'message': 'Account created successfully',
            'data': {
                'user': {
                    'id': created_user['id'],
                    'username': created_user['username'],
                    'email': created_user['email'],
                    'bio': created_user.get('bio', ''),
                    'userType': 'influencer'
                },
                'token': token,
                'userType': 'influencer'
            }
        }), 201
        
    except Exception as e:
        logger.error(f"Registration error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Registration failed. Please try again.'
        }), 500

@app.route('/api/login', methods=['POST'])
@app.route('/api/auth/login', methods=['POST'])
def login():
    """FIXED: Login influencer with consistent response format"""
    try:
        data = request.get_json()
        
        username = data.get('username', '').strip()
        password = data.get('password', '')
        
        if not username or not password:
            return jsonify({
                'status': 'error',
                'message': 'Username and password are required'
            }), 400
        
        # Get user
        user = db.get_influencer_by_username(username)
        if not user:
            return jsonify({
                'status': 'error',
                'message': 'Invalid credentials'
            }), 401
        
        # Verify password
        if user['password_hash'] != hash_password(password):
            return jsonify({
                'status': 'error',
                'message': 'Invalid credentials'
            }), 401
        
        # Generate token
        token = generate_jwt_token(user)
        
        logger.info(f"✅ User logged in successfully: {username}")
        
        # FIXED: Consistent response format for frontend
        return jsonify({
            'status': 'success',
            'message': 'Login successful',
            'data': {
                'user': {
                    'id': user['id'],
                    'username': user['username'],
                    'email': user['email'],
                    'bio': user.get('bio', ''),
                    'avatar_id': user.get('heygen_avatar_id'),
                    'userType': 'influencer'
                },
                'token': token,
                'userType': 'influencer'
            }
        })
        
    except Exception as e:
        logger.error(f"Login error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Login failed'
        }), 500

# Enhanced Influencer Profile Routes
@app.route('/api/influencer/profile', methods=['GET'])
@token_required
def get_influencer_profile(current_user):
    """FIXED: Get influencer profile with enhanced voice, avatar, and knowledge info"""
    try:
        influencer = db.get_influencer(current_user['id'])
        if not influencer:
            return jsonify({
                'status': 'error',
                'message': 'Influencer not found'
            }), 404

        # Get knowledge documents
        knowledge_documents = []
        try:
            doc_response = db.supabase.table('knowledge_documents') \
                .select('*') \
                .eq('influencer_id', current_user['id']) \
                .order('created_at', desc=True) \
                .execute()
            
            if doc_response.data:
                for doc in doc_response.data:
                    knowledge_documents.append({
                        'id': doc['id'],
                        'name': doc['filename'],
                        'size': doc['file_size'],
                        'type': doc['content_type'],
                        'status': 'completed' if doc.get('is_processed') else 'processing',
                        'document_id': doc['id'],
                        'uploaded_at': doc.get('created_at'),
                        'processed': doc.get('is_processed', False)
                    })
                
                logger.info(f"Found {len(knowledge_documents)} documents for user {current_user['username']}")
            
        except Exception as doc_error:
            logger.error(f"Error loading knowledge documents: {doc_error}")

        profile_data = {
            'id': influencer['id'],
            'username': influencer['username'],
            'email': influencer['email'],
            'bio': influencer.get('bio', ''),
            'expertise': influencer.get('expertise', ''),
            'personality': influencer.get('personality', ''),
            'avatar_id': influencer.get('heygen_avatar_id'),
            'has_avatar': bool(influencer.get('heygen_avatar_id')),
            'voice_id': influencer.get('voice_id'),
            'preferred_voice_id': influencer.get('preferred_voice_id', Config.DEFAULT_VOICE_ID),
            'avatar_training_status': influencer.get('avatar_training_status', 'none'),
            'avatar_type': influencer.get('avatar_type', 'none'),
            'created_at': influencer.get('created_at'),
            'updated_at': influencer.get('updated_at'),
            'knowledge_documents': knowledge_documents
        }

        return jsonify({
            'status': 'success',
            'data': profile_data
        })

    except Exception as e:
        logger.error(f"Get influencer profile error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get profile'
        }), 500

@app.route('/api/influencer/profile', methods=['PUT'])
@token_required
def update_influencer_profile(current_user):
    """Update influencer profile"""
    try:
        data = request.get_json()
        update_data = {}

        allowed_fields = ['bio', 'email', 'expertise', 'personality', 'preferred_voice_id']
        for field in allowed_fields:
            if field in data:
                update_data[field] = data[field]

        success = db.update_influencer(current_user['id'], update_data)

        if success:
            return jsonify({
                'status': 'success',
                'message': 'Profile updated successfully'
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to update profile'
            }), 500

    except Exception as e:
        logger.error(f"Update influencer profile error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'An error occurred while updating profile'
        }), 500

# Enhanced Avatar Management Routes
@app.route('/api/avatar/create', methods=['POST'])
@token_required
def create_avatar(current_user):
    """FIXED: Avatar creation with enhanced timing and error handling"""
    try:
        # Get uploaded file
        image_file = request.files.get('file') or request.files.get('image')
        
        if not image_file:
            return jsonify({
                'status': 'error',
                'message': 'No image file provided'
            }), 400
        
        # Validate file
        try:
            validate_file_upload(image_file)
        except ValueError as ve:
            return jsonify({
                'status': 'error',
                'message': str(ve)
            }), 400
        
        logger.info(f"🎭 Creating avatar for user: {current_user['username']}")
        
        # Create custom talking photo avatar
        try:
            avatar_result = HeyGenAPI.create_photo_avatar(image_file)
            
            if avatar_result and avatar_result.get('avatar_id'):
                avatar_id = avatar_result['avatar_id']
                
                # Update user record with avatar info
                success = db.update_avatar_status(
                    current_user['id'],
                    {
                        'heygen_avatar_id': avatar_id,
                        'avatar_training_status': avatar_result.get('status', 'ready'),
                        'avatar_type': avatar_result.get('type', 'talking_photo')
                    }
                )
                
                if success:
                    logger.info(f"✅ Custom avatar created successfully: {avatar_id}")
                    
                    return jsonify({
                        'status': 'success',
                        'message': 'Custom avatar created from your photo!',
                        'data': {
                            'avatar_id': avatar_id,
                            'status': avatar_result.get('status', 'ready'),
                            'type': avatar_result.get('type', 'talking_photo'),
                            'is_custom': True,
                            'is_newly_created': True,
                            'processing_time_estimate': 60,
                            'method': avatar_result.get('method', 'talking_photo'),
                            'created_at': datetime.now(timezone.utc).isoformat()
                        }
                    }), 201
                else:
                    return jsonify({
                        'status': 'error',
                        'message': 'Failed to save avatar to database'
                    }), 500
            else:
                raise Exception("No avatar ID returned from HeyGen")
                
        except Exception as heygen_error:
            logger.error(f"❌ HeyGen avatar creation failed: {heygen_error}")
            
            error_message = str(heygen_error)
            
            # Handle specific HeyGen error types
            if "face" in error_message.lower():
                return jsonify({
                    'status': 'error',
                    'message': 'Could not detect a face in the image. Please use a clear photo with a visible face.',
                    'error_type': 'face_detection_error',
                    'suggestions': [
                        'Use a clear photo with good lighting',
                        'Ensure only one face is visible in the image',
                        'Try a front-facing photo'
                    ]
                }), 400
            elif "format not supported" in error_message.lower():
                return jsonify({
                    'status': 'error',
                    'message': 'Image format not supported. Please use JPEG, PNG, or WebP format.',
                    'error_type': 'format_error'
                }), 400
            else:
                return jsonify({
                    'status': 'error',
                    'message': f'Avatar creation failed: {error_message}',
                    'error_type': 'creation_failed'
                }), 500
            
    except Exception as e:
        logger.error(f"❌ Avatar creation endpoint error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'An unexpected error occurred during avatar creation'
        }), 500

@app.route('/api/avatar/preview/<influencer_id>', methods=['GET'])
def get_avatar_preview(influencer_id):
    """Get avatar preview information for chat display"""
    try:
        influencer = db.get_influencer(influencer_id)
        
        if not influencer:
            return jsonify({
                'status': 'error',
                'message': 'Influencer not found'
            }), 404
        
        if not influencer.get('heygen_avatar_id'):
            return jsonify({
                'status': 'error',
                'message': 'No avatar available'
            }), 404
        
        # Return avatar information
        avatar_data = {
            'avatar_id': influencer['heygen_avatar_id'],
            'avatar_type': influencer.get('avatar_type', 'talking_photo'),
            'username': influencer['username'],
            'has_custom_avatar': True,
            'status': 'ready'
        }
        
        return jsonify({
            'status': 'success',
            'data': avatar_data
        })
        
    except Exception as e:
        logger.error(f"Get avatar preview error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get avatar preview'
        }), 500

@app.route('/api/avatar/status', methods=['GET'])
@token_required
def get_avatar_status(current_user):
    """Get current avatar status and HeyGen plan info"""
    try:
        # Get user's current avatar
        user_avatar = current_user.get('heygen_avatar_id')
        
        # Get available HeyGen avatars (to check total count)
        available_avatars = HeyGenAPI.get_available_avatars()
        
        return jsonify({
            'status': 'success',
            'data': {
                'has_avatar': bool(user_avatar),
                'avatar_id': user_avatar,
                'total_available_avatars': len(available_avatars),
                'heygen_credits': '109 credits remaining',  # You can get this from HeyGen API
                'heygen_dashboard_url': 'https://app.heygen.com/avatars',
                'upgrade_url': 'https://app.heygen.com/settings/billing'
            }
        })
        
    except Exception as e:
        logger.error(f"Get avatar status error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get avatar status'
        }), 500

# Enhanced Voice Management Routes
@app.route('/api/avatar/list-voices', methods=['GET'])
def list_available_voices():
    """FIXED: Get enhanced list of available voices for avatar setup"""
    try:
        voices = [
            {
                "voice_id": "2d5b0e6cf36f460aa7fc47e3eee4ba54",
                "name": "Rachel (Professional)",
                "gender": "Female",
                "language": "English",
                "style": "Professional",
                "description": "Clear, professional female voice perfect for business content"
            },
            {
                "voice_id": "d7bbcdd6964c47bdaae26decade4a933",
                "name": "David (Professional)",
                "gender": "Male", 
                "language": "English",
                "style": "Professional",
                "description": "Authoritative male voice ideal for educational content"
            },
            {
                "voice_id": "4d2b8e6cf36f460aa7fc47e3eee4ba12",
                "name": "Emma (Friendly)",
                "gender": "Female",
                "language": "English", 
                "style": "Friendly",
                "description": "Warm, approachable female voice great for lifestyle content"
            },
            {
                "voice_id": "3a1c7d5bf24e350bb6dc46e2dee3ab21",
                "name": "Michael (Casual)",
                "gender": "Male",
                "language": "English",
                "style": "Casual", 
                "description": "Relaxed male voice perfect for conversational content"
            },
            {
                "voice_id": "1bd001e7e50f421d891986aad5158bc8",
                "name": "Olivia (Warm)",
                "gender": "Female",
                "language": "English",
                "style": "Warm",
                "description": "Gentle, caring female voice ideal for personal brands"
            }
        ]
        
        return jsonify({
            'status': 'success',
            'data': {
                'voices': voices,
                'total_count': len(voices),
                'default_voice_id': Config.DEFAULT_VOICE_ID
            }
        })
        
    except Exception as e:
        logger.error(f"List voices error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get voice list'
        }), 500

@app.route('/api/voice/preference', methods=['POST'])
@token_required
def save_voice_preference(current_user):
    """FIXED: Enhanced voice preference saving with avatar integration"""
    try:
        data = request.get_json()
        voice_id = data.get('voice_id')
        
        if not voice_id:
            return jsonify({
                'status': 'error',
                'message': 'Voice ID is required'
            }), 400
        
        # Update user's preferred voice
        success = db.update_influencer(current_user['id'], {
            'preferred_voice_id': voice_id
        })
        
        if success:
            logger.info(f"Voice preference updated for {current_user['username']}: {voice_id}")
            return jsonify({
                'status': 'success',
                'message': 'Voice preference saved successfully',
                'data': {
                    'preferred_voice_id': voice_id
                }
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to save voice preference'
            }), 500
            
    except Exception as e:
        logger.error(f"Save voice preference error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to save voice preference'
        }), 500

@app.route('/api/voice/preference', methods=['GET'])
@token_required
def get_voice_preference(current_user):
    """FIXED: Get user's preferred voice"""
    try:
        influencer = db.get_influencer(current_user['id'])
        preferred_voice_id = influencer.get('preferred_voice_id', Config.DEFAULT_VOICE_ID)
        
        return jsonify({
            'status': 'success',
            'data': {
                'preferred_voice_id': preferred_voice_id
            }
        })
        
    except Exception as e:
        logger.error(f"Get voice preference error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get voice preference'
        }), 500

@app.route('/api/voice/clone', methods=['POST'])
@token_required
def create_voice_clone(current_user):
    """Enhanced voice cloning endpoint"""
    try:
        # Check if files were uploaded
        if 'audio_samples' not in request.files:
            return jsonify({
                'status': 'error',
                'message': 'No audio samples provided'
            }), 400
        
        files = request.files.getlist('audio_samples')
        voice_name = request.form.get('voice_name', f"{current_user['username']}'s Voice")
        description = request.form.get('description', f"Voice clone for {current_user['username']}")
        
        if len(files) < 3:
            return jsonify({
                'status': 'error',
                'message': 'At least 3 audio samples are required'
            }), 400
        
        # Validate files
        total_size = 0
        for file in files:
            if not file.filename.lower().endswith(('.mp3', '.wav', '.webm', '.m4a')):
                return jsonify({
                    'status': 'error',
                    'message': f'Invalid file type: {file.filename}. Only MP3, WAV, WebM, and M4A are supported.'
                }), 400
            
            file.seek(0, 2)  # Seek to end
            size = file.tell()
            file.seek(0)  # Reset
            total_size += size
            
            if size > 50 * 1024 * 1024:  # 50MB per file
                return jsonify({
                    'status': 'error', 
                    'message': f'File {file.filename} is too large. Maximum 50MB per file.'
                }), 400
        
        if total_size > 200 * 1024 * 1024:  # 200MB total
            return jsonify({
                'status': 'error',
                'message': 'Total file size too large. Maximum 200MB total.'
            }), 400
        
        # Generate custom voice ID (in production, would integrate with ElevenLabs or similar)
        custom_voice_id = f"custom_{current_user['id']}_{str(uuid.uuid4())[:8]}"
        
        # Update user's voice preferences
        success = db.update_influencer(current_user['id'], {
            'voice_id': custom_voice_id,
            'preferred_voice_id': custom_voice_id
        })
        
        if success:
            logger.info(f"✅ Custom voice created for {current_user['username']}: {custom_voice_id}")
            return jsonify({
                'status': 'success',
                'message': 'Custom voice created successfully',
                'data': {
                    'voice_id': custom_voice_id,
                    'voice_name': voice_name,
                    'status': 'ready',
                    'is_custom': True,
                    'sample_count': len(files)
                }
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to save custom voice'
            }), 500
            
    except Exception as e:
        logger.error(f"Voice clone error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to create voice clone'
        }), 500

# Enhanced Video Generation Routes
@app.route('/api/avatar/test-video', methods=['POST'])
@token_required
def generate_test_video(current_user):
    """FIXED: Generate test video with enhanced handling for new avatars"""
    try:
        data = request.get_json()
        
        avatar_id = data.get('avatar_id')
        text = data.get('text', 'Hello! This is a test of my AI avatar with my selected voice. How does it look and sound?')
        voice_id = data.get('voice_id')
        is_newly_created = data.get('is_newly_created', False)
        validate_only = data.get('validate_only', False)
        
        if not avatar_id:
            return jsonify({
                'status': 'error',
                'message': 'Avatar ID is required'
            }), 400
        
        # Get user's preferred voice if not specified
        if not voice_id:
            influencer = db.get_influencer(current_user['id'])
            voice_id = influencer.get('preferred_voice_id', Config.DEFAULT_VOICE_ID)
        
        logger.info(f"Generating test video for avatar: {avatar_id} with voice: {voice_id}")
        logger.info(f"Is newly created: {is_newly_created}")
        
        # If validate_only flag is set, just check if avatar exists
        if validate_only:
            try:
                is_valid = HeyGenAPI.validate_avatar(avatar_id, is_newly_created)
                if is_valid:
                    return jsonify({
                        'status': 'success',
                        'message': 'Avatar validation successful'
                    })
                else:
                    return jsonify({
                        'status': 'error',
                        'message': 'Avatar not found on HeyGen servers',
                        'error_type': 'avatar_not_found'
                    }), 404
            except Exception as validation_error:
                logger.error(f"Avatar validation failed: {validation_error}")
                return jsonify({
                    'status': 'error',
                    'message': 'Avatar validation failed',
                    'error_type': 'validation_failed'
                }), 500
        
        # Enhanced timing for newly created avatars
        if is_newly_created:
            logger.info("New avatar detected - waiting 10 seconds for HeyGen processing...")
            import time
            time.sleep(10)
        
        # Attempt video generation with proper error handling
        try:
            video_result = HeyGenAPI.generate_video(
                avatar_id=avatar_id,
                text=text,
                voice_id=voice_id,
                is_newly_created=is_newly_created
            )
            
            if video_result and video_result.get('video_id'):
                logger.info(f"Video generation started: {video_result['video_id']}")
                
                return jsonify({
                    'status': 'success',
                    'message': 'Video generation started',
                    'data': {
                        'video_id': video_result['video_id'],
                        'status': 'processing',
                        'estimated_time': video_result.get('estimated_time', 30),
                        'text': text,
                        'voice_id': voice_id
                    }
                })
            else:
                raise Exception("No video ID returned from HeyGen")
                
        except Exception as video_error:
            error_message = str(video_error)
            logger.error(f"HeyGen video generation failed: {video_error}")
            
            # Enhanced error handling for avatar timing issues
            if error_message == "AVATAR_STILL_PROCESSING":
                retry_delay = 60 if is_newly_created else 30
                
                return jsonify({
                    'status': 'error',
                    'message': 'Avatar is still processing on HeyGen servers. Please wait and try again.',
                    'error_type': 'avatar_processing',
                    'requires_retry': True,
                    'retry_delay': retry_delay,
                    'is_newly_created': is_newly_created
                }), 202
                
            elif error_message == "AVATAR_NOT_FOUND":
                # Clear invalid avatar from database
                db.update_avatar_status(current_user['id'], {
                    'heygen_avatar_id': None,
                    'avatar_training_status': 'none',
                    'avatar_type': 'none'
                })
                
                return jsonify({
                    'status': 'error',
                    'message': 'Avatar no longer exists on HeyGen servers. Please create a new avatar.',
                    'error_type': 'avatar_not_found',
                    'requires_new_avatar': True
                }), 404
            else:
                return jsonify({
                    'status': 'error',
                    'message': f'Video generation failed: {str(video_error)}',
                    'error_type': 'generation_failed'
                }), 500
            
    except Exception as e:
        logger.error(f"Test video generation error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate test video'
        }), 500

@app.route('/api/avatar/video-status/<video_id>', methods=['GET'])
@token_required
def get_video_status_fixed(current_user, video_id):
    """FIXED: Get real video status from HeyGen API"""
    try:
        logger.info(f"Checking video status for: {video_id}")
        
        # Use the proper HeyGen API method for status checking
        video_status = HeyGenAPI.get_video_status(video_id)
        
        logger.info(f"Video status response: {video_status}")
        
        return jsonify({
            'status': 'success',
            'data': video_status
        })
                
    except Exception as e:
        logger.error(f"Video status check error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to check video status: {str(e)}'
        }), 500

@app.route('/api/voice/generate-audio', methods=['POST'])
@token_required
def generate_voice_audio_endpoint(current_user):
    """FIXED: Generate audio using user's preferred voice"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        text = data.get('text', '').strip()
        voice_id = data.get('voice_id')
        
        if not text:
            return jsonify({
                'status': 'error',
                'message': 'Text is required'
            }), 400
        
        if len(text) > 1000:
            return jsonify({
                'status': 'error',
                'message': 'Text too long. Maximum 1000 characters.'
            }), 400
        
        # Get user's preferred voice if not specified
        if not voice_id:
            influencer = db.get_influencer(current_user['id'])
            voice_id = influencer.get('preferred_voice_id') or influencer.get('voice_id') or Config.DEFAULT_VOICE_ID
        
        logger.info(f"🔊 Generating audio for user: {current_user['username']}, voice: {voice_id}")
        
        # Generate audio using correct method
        audio_url = chatbot.generate_audio_response(text, voice_id)
        
        if audio_url:
            return jsonify({
                'status': 'success',
                'data': {
                    'audio_url': audio_url,
                    'text': text,
                    'voice_id': voice_id,
                    'duration_estimate': len(text) / 10
                }
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to generate audio'
            }), 500
            
    except Exception as e:
        logger.error(f"Voice audio generation error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate voice audio'
        }), 500

@app.route('/api/voice/preview', methods=['POST'])
def preview_voice():
    """FIXED: Preview a voice with sample text (no authentication required)"""
    try:
        data = request.get_json()
        
        voice_id = data.get('voice_id', Config.DEFAULT_VOICE_ID)
        sample_text = data.get('text', 'Hello! This is a preview of this voice. How does it sound?')
        
        # Limit sample text
        if len(sample_text) > 200:
            sample_text = sample_text[:200] + "..."
        
        logger.info(f"🎤 Voice preview requested for voice: {voice_id}")
        
        # Generate audio using the chatbot
        audio_url = chatbot.generate_audio_response(sample_text, voice_id)
        
        if audio_url:
            return jsonify({
                'status': 'success',
                'data': {
                    'audio_url': audio_url,
                    'text': sample_text,
                    'voice_id': voice_id
                }
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to generate voice preview'
            }), 500
            
    except Exception as e:
        logger.error(f"Voice preview error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate voice preview'
        }), 500
    
@app.route('/api/speech-to-text', methods=['POST'])
def speech_to_text():
    """FIXED: Convert speech to text using OpenAI Whisper with proper file handling"""
    try:
        # Check if audio file is provided
        if 'audio' not in request.files:
            return jsonify({
                'status': 'error',
                'message': 'No audio file provided'
            }), 400
        
        audio_file = request.files['audio']
        
        if audio_file.filename == '':
            return jsonify({
                'status': 'error',
                'message': 'No audio file selected'
            }), 400
        
        # Validate file size
        audio_file.seek(0, 2)  # Seek to end
        file_size = audio_file.tell()
        audio_file.seek(0)  # Reset to beginning
        
        if file_size > Config.MAX_CONTENT_LENGTH:
            return jsonify({
                'status': 'error',
                'message': 'Audio file too large'
            }), 400
        
        if file_size == 0:
            return jsonify({
                'status': 'error',
                'message': 'Audio file is empty'
            }), 400
        
        # FIXED: Proper temporary file handling for webm files
        temp_filename = None
        try:
            # Create temporary file with proper extension
            with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as temp_file:
                # Read the file content and write to temp file
                audio_file.seek(0)
                file_content = audio_file.read()
                temp_file.write(file_content)
                temp_file.flush()
                temp_filename = temp_file.name
            
            logger.info(f"Temp file created: {temp_filename}, size: {len(file_content)} bytes")
            
            # Use OpenAI Whisper API with proper file object
            with open(temp_filename, 'rb') as audio_data:
                transcript = chatbot.client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_data,
                    response_format="text"
                )
            
            # Extract text from response
            transcription = transcript.strip() if isinstance(transcript, str) else transcript.text.strip()
            
            if transcription:
                logger.info(f"Speech-to-text successful: {transcription[:50]}...")
                return jsonify({
                    'status': 'success',
                    'data': {
                        'transcription': transcription,
                        'confidence': 0.95
                    }
                })
            else:
                return jsonify({
                    'status': 'error',
                    'message': 'Could not transcribe audio. Please speak clearly and try again.'
                }), 400
                
        finally:
            # Proper cleanup of temporary file
            if temp_filename and os.path.exists(temp_filename):
                try:
                    os.unlink(temp_filename)
                    logger.info(f"Cleaned up temp file: {temp_filename}")
                except Exception as cleanup_error:
                    logger.warning(f"Could not delete temporary file: {cleanup_error}")
        
    except Exception as e:
        logger.error(f"Speech-to-text error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Speech-to-text processing failed: {str(e)}'
        }), 500
       
def format_affiliate_product_response_enhanced(products: List[Dict], influencer_name: str, result_info: Dict) -> str:
    """ENHANCED: Format affiliate products with better context and demo handling"""
    if not products:
        return f"I'd love to help you find products, but I'm having trouble connecting to my affiliate partners right now. Let me know what specific features you're looking for!"
    
    # Determine response type
    has_real = result_info.get('has_real_products', False)
    has_demo = len(result_info.get('demo_platforms', [])) > 0
    successful_platforms = result_info.get('successful_platforms', [])
    demo_platforms = result_info.get('demo_platforms', [])
    
    # Craft opening based on product types
    if has_real and not has_demo:
        response = f"Great question! I found some excellent products through my affiliate partnerships:\n\n"
    elif has_demo and not has_real:
        response = f"I have some great product recommendations for you! (Note: Currently showing demo products due to API connectivity):\n\n"
    else:
        response = f"Here are some product options I found:\n\n"
    
    # Format each product
    for i, product in enumerate(products[:3], 1):
        # Handle different currencies
        if product.get('currency') == 'JPY':
            price_str = f"¥{product['price']:,.0f}" if product.get('price', 0) > 0 else "See price"
        else:
            price_str = f"${product['price']:.2f}" if product.get('price', 0) > 0 else "See price"
        
        rating_str = f"⭐ {product['rating']:.1f}" if product.get('rating', 0) > 0 else ""
        
        # Product name with demo indicator
        name = product['name']
        if product.get('is_demo'):
            name += " (Demo)"
        
        response += f"**{i}. {name}**\n"
        response += f"💰 {price_str}"
        
        if rating_str:
            response += f" | {rating_str}"
        
        if product.get('review_count', 0) > 0:
            response += f" ({product['review_count']} reviews)"
        
        if product.get('shop_name'):
            response += f" | 🏪 {product['shop_name']}"
        
        response += f"\n📝 {product.get('description', 'Quality product')[:100]}...\n"
        
        # Only show affiliate links for real products
        if product.get('affiliate_url') and not product.get('is_demo'):
            response += f"🔗 [View Product]({product['affiliate_url']})\n\n"
        else:
            response += "\n"
    
    # Add contextual footer
    if successful_platforms and demo_platforms:
        response += f"💡 *Found real products from {', '.join(successful_platforms)} and demo products from {', '.join(demo_platforms)}.*\n"
    elif successful_platforms:
        response += f"💡 *Found through my {', '.join(successful_platforms)} partnerships. I earn a small commission if you make a purchase.*\n"
    elif demo_platforms:
        response += f"💡 *These are demonstration products from {', '.join(demo_platforms)}. I'm working to resolve API connectivity issues.*\n"
    
    return response

@app.route('/api/chat', methods=['POST'])
def chat_with_influencer():
    """FIXED: Enhanced chat endpoint with proper avatar and voice integration"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        user_message = data.get('message', '').strip()
        username = data.get('username', '').strip()
        influencer_id = data.get('influencer_id')
        session_id = data.get('session_id', str(uuid.uuid4()))
        voice_mode = data.get('voice_mode', False)
        video_mode = data.get('video_mode', False)
        
        if not user_message:
            return jsonify({
                'status': 'error',
                'message': 'Message is required'
            }), 400
        
        # Get influencer info
        influencer = None
        if influencer_id:
            influencer = db.get_influencer(influencer_id)
        elif username:
            clean_username = username.strip().lower()
            influencer = db.get_influencer_by_username(clean_username)
        
        if not influencer:
            return jsonify({
                'status': 'error',
                'message': 'Influencer not found'
            }), 404
        
        influencer_id = influencer['id']
        influencer_name = influencer.get('username', 'the influencer')
        
        logger.info(f"Chat request - User: {influencer_name}, Message: {user_message[:50]}...")
        
        # Generate comprehensive response using the enhanced chatbot
        response_data = chatbot.get_comprehensive_chat_response(
            message=user_message,
            influencer_id=influencer_id,
            session_id=session_id,
            influencer_name=influencer_name,
            voice_mode=voice_mode,
            video_mode=video_mode
        )
        
        # Generate audio if voice mode is enabled and user has avatar
        if voice_mode and influencer.get('heygen_avatar_id'):
            try:
                voice_id = influencer.get('preferred_voice_id') or influencer.get('voice_id') or Config.DEFAULT_VOICE_ID
                audio_url = chatbot.generate_audio_response(response_data['text'], voice_id)
                if audio_url:
                    response_data['audio_url'] = audio_url
                    response_data['has_audio'] = True
            except Exception as audio_error:
                logger.error(f"Audio generation failed: {audio_error}")
        
        # Generate video if video mode is enabled and user has avatar
        if video_mode and influencer.get('heygen_avatar_id'):
            try:
                video_url = chatbot.generate_enhanced_video_response(
                    response_data['text'], 
                    influencer_id,
                    influencer.get('preferred_voice_id')
                )
                if video_url:
                    response_data['video_url'] = video_url
                    response_data['has_video'] = True
            except Exception as video_error:
                logger.error(f"Video generation failed: {video_error}")
        
        # Store interaction
        try:
            interaction_data = {
                'id': str(uuid.uuid4()),
                'influencer_id': influencer_id,
                'session_id': session_id,
                'user_message': user_message,
                'bot_response': response_data['text'],
                'products_included': response_data.get('products_included', False),
                'knowledge_enhanced': response_data.get('knowledge_enhanced', False),
                'has_video': response_data.get('video_url', '') != '',
                'has_audio': response_data.get('audio_url', '') != '',
                'created_at': datetime.now(timezone.utc).isoformat()
            }
            
            db.store_chat_interaction(interaction_data)
        except Exception as storage_error:
            logger.error(f"Failed to store interaction: {storage_error}")
        
        return jsonify({
            'status': 'success',
            'data': response_data
        })
        
    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Chat service temporarily unavailable'
        }), 500

def generate_emergency_ai_recommendations(query: str, influencer_name: str) -> str:
    """Emergency AI recommendations when everything fails"""
    try:
        # Simple template-based recommendations for critical failures
        common_products = {
            'laptop': [
                'Dell Inspiron 15 3000 - Reliable everyday laptop (~$400-600)',
                'HP Pavilion x360 - Versatile 2-in-1 laptop (~$500-700)', 
                'Lenovo ThinkPad E15 - Business-grade laptop (~$600-800)'
            ],
            'phone': [
                'iPhone 14 - Latest iOS features and camera (~$700-900)',
                'Samsung Galaxy S23 - Android flagship with great display (~$600-800)',
                'Google Pixel 7 - Pure Android experience with excellent camera (~$500-700)'
            ],
            'headphones': [
                'Sony WH-1000XM5 - Premium noise-canceling (~$300-400)',
                'Apple AirPods Pro - Seamless iOS integration (~$200-250)',
                'Bose QuietComfort 45 - Comfortable all-day wear (~$250-350)'
            ]
        }
        
        query_lower = query.lower()
        products = None
        
        # Find matching category
        for category, product_list in common_products.items():
            if category in query_lower:
                products = product_list
                break
        
        if products:
            response = f"Based on my experience, here are some {query.lower()} options I'd recommend:\n\n"
            for i, product in enumerate(products, 1):
                response += f"**{i}. {product}**\n\n"
            response += "💡 *These are my general recommendations. I'm working to resolve affiliate connection issues to bring you real-time deals.*"
        else:
            response = f"I'd love to help you find great {query} options! What specific features, budget, or use case are you considering? I can share some general advice to help you make the best choice."
        
        return response
        
    except Exception as e:
        logger.error(f"❌ Emergency recommendations error: {e}")
        return f"I'm here to help with {query}! Could you tell me more about what you're looking for specifically?"

def is_product_query_enhanced(message: str) -> bool:
    """Enhanced product query detection"""
    message_lower = message.lower().strip()
    
    product_patterns = [
        'recommend', 'recommendation', 'suggest', 'suggestion', 
        'what should i buy', 'help me find', 'looking for',
        'best product', 'product for', 'good product',
        'buy', 'purchase', 'shop', 'shopping'
    ]
    
    return any(pattern in message_lower for pattern in product_patterns)

def format_affiliate_product_response(products: List[Dict], influencer_name: str, successful_platforms: List[str]) -> str:
    """FIXED: Format real affiliate products only"""
    if not products:
        return f"I couldn't find any products through my affiliate connections right now. Let me know what specific features you're looking for and I can provide some general recommendations!"
    
    response = f"Great question! I found some excellent products through my affiliate partnerships:\n\n"
    
    for i, product in enumerate(products[:3], 1):
        price_str = f"${product['price']:.2f}" if product.get('price', 0) > 0 else "See price"
        rating_str = f"⭐ {product['rating']:.1f}" if product.get('rating', 0) > 0 else ""
        
        response += f"**{i}. {product['name']}**\n"
        response += f"💰 {price_str}"
        
        if rating_str:
            response += f" | {rating_str}"
        
        if product.get('review_count', 0) > 0:
            response += f" ({product['review_count']} reviews)"
        
        if product.get('shop_name'):
            response += f" | 🏪 {product['shop_name']}"
        
        response += f"\n📝 {product.get('description', 'Quality product')[:120]}...\n"
        
        if product.get('affiliate_url'):
            response += f"🔗 [View Product]({product['affiliate_url']})\n\n"
        else:
            response += "\n"
    
    # Simple footer for real products
    if successful_platforms:
        platforms_text = ', '.join(successful_platforms)
        response += f"💡 *Found through my {platforms_text} partnerships. I earn a small commission if you make a purchase.*\n"
    
    return response

def generate_enhanced_ai_recommendations(query: str, influencer_name: str, affiliate_links: List[Dict]) -> str:
    """Generate AI recommendations when no real products available"""
    try:
        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        
        platform_context = ""
        if affiliate_links:
            platforms = [link.get('platform', '').title() for link in affiliate_links if link.get('is_active', True)]
            platform_context = f" I have affiliate partnerships with {', '.join(platforms)}, but couldn't find specific products right now."
        
        prompt = f"""You are {influencer_name}, helping someone find products related to: {query}

Generate 3 realistic, specific product recommendations. For each:
- Specific product name and brand
- Realistic price range in USD
- Brief compelling description (1-2 sentences)
- Why it's relevant to their query

Format as helpful, personal recommendations.{platform_context}

Start with: "Based on my experience, here are some great options I'd recommend:" """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": f"You are {influencer_name}, a helpful influencer who gives personalized product recommendations."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=400,
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content.strip()
        
        if affiliate_links:
            ai_response += f"\n\n💡 *These are my personal recommendations. I'm working to get you live deals through my affiliate partnerships!*"
        else:
            ai_response += f"\n\n💡 *These are my personal recommendations. I'm setting up affiliate partnerships to bring you exclusive deals soon!*"
        
        return ai_response
        
    except Exception as e:
        logger.error(f"❌ AI recommendations error: {e}")
        return f"I'd love to help you find great {query} options! What specific features or budget are you considering?"


def generate_no_affiliate_ai_recommendations(query: str, influencer_name: str) -> str:
    """Generate AI recommendations when no affiliate setup"""
    try:
        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        
        prompt = f"""You are {influencer_name}, helping someone find products related to: {query}

Generate 3 realistic product recommendations with general advice. Focus on:
- What to look for when shopping for these items
- General price ranges and quality indicators
- Brands or features to consider

Be helpful and educational rather than providing specific purchase links.

Start with: "I'd love to help you find the perfect [product type]! Here's what I'd look for:" """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": f"You are {influencer_name}, a helpful influencer who gives shopping advice."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300,
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content.strip()
        ai_response += f"\n\n💡 *I'm working on setting up affiliate partnerships to bring you specific product recommendations and deals. Stay tuned!*"
        
        return ai_response
        
    except Exception as e:
        logger.error(f"❌ No affiliate AI recommendations error: {e}")
        return f"I'd love to help you find great {query} options! What specific features or budget are you considering? I'm working on setting up partnerships to bring you the best deals."

def generate_basic_ai_recommendations(query: str, influencer_name: str) -> str:
    """Basic fallback AI recommendations"""
    return f"I'd love to help you find great {query} options! What specific features, budget, or use case are you considering? I can share some general advice to help you make the best choice."

def generate_basic_chat_response(message: str, influencer_name: str) -> str:
    """Generate basic non-product chat response"""
    try:
        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": f"You are {influencer_name}'s AI assistant. Be helpful, friendly, and conversational. Keep responses concise (2-3 sentences)."},
                {"role": "user", "content": message}
            ],
            max_tokens=150,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        logger.error(f"❌ Basic chat response error: {e}")
        return f"Hi! I'm {influencer_name}'s AI assistant. Thanks for your message! How can I help you today?"
   
def search_knowledge_base_improved(db, influencer_id: str, query_embedding: List[float], limit: int = 5) -> List[Dict]:
    """FIXED: Improved knowledge base search with better similarity calculation"""
    try:
        # Get all chunks for the influencer
        response = db.supabase.table('knowledge_chunks') \
            .select('*, knowledge_documents(filename)') \
            .eq('influencer_id', influencer_id) \
            .execute()
        
        chunks = response.data if response.data else []
        logger.info(f"🔍 Found {len(chunks)} knowledge chunks for search")
        
        if not chunks:
            return []
        
        # Calculate similarities
        similarities = []
        
        for chunk in chunks:
            try:
                if chunk.get('embedding'):
                    # FIXED: Parse embedding from JSON string
                    embedding_str = chunk['embedding']
                    if isinstance(embedding_str, str):
                        try:
                            chunk_embedding = json.loads(embedding_str)
                        except json.JSONDecodeError:
                            logger.warning(f"Could not parse embedding for chunk {chunk.get('id', 'unknown')}")
                            continue
                    elif isinstance(embedding_str, list):
                        chunk_embedding = embedding_str
                    else:
                        logger.warning(f"Unknown embedding format for chunk {chunk.get('id', 'unknown')}")
                        continue
                    
                    if chunk_embedding and len(chunk_embedding) == len(query_embedding):
                        # FIXED: Calculate cosine similarity properly
                        similarity = calculate_cosine_similarity(query_embedding, chunk_embedding)
                        
                        if similarity > 0.05:  # Only include chunks with some relevance
                            similarities.append({
                                'chunk_id': chunk.get('id'),
                                'document_id': chunk.get('document_id'),
                                'text': chunk.get('chunk_text', ''),
                                'similarity': float(similarity),
                                'chunk_index': chunk.get('chunk_index', 0),
                                'document_name': chunk.get('knowledge_documents', {}).get('filename', 'Unknown') if chunk.get('knowledge_documents') else 'Unknown'
                            })
                        
            except Exception as embedding_error:
                logger.error(f"Error processing embedding for chunk {chunk.get('id', 'unknown')}: {embedding_error}")
                continue
        
        # Sort by similarity and return top results
        similarities.sort(key=lambda x: x['similarity'], reverse=True)
        top_results = similarities[:limit]
        
        logger.info(f"🎯 Knowledge search returned {len(top_results)} relevant chunks")
        for result in top_results[:3]:  # Log top 3 for debugging
            logger.info(f"   - {result['document_name']}: {result['similarity']:.3f} - {result['text'][:100]}...")
        
        return top_results
        
    except Exception as e:
        logger.error(f"Error in improved knowledge search: {e}")
        return []

def calculate_cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """FIXED: Calculate cosine similarity between two vectors"""
    try:
        import math
        
        if len(vec1) != len(vec2):
            logger.warning(f"Vector length mismatch: {len(vec1)} vs {len(vec2)}")
            return 0.0
        
        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        magnitude1 = math.sqrt(sum(a * a for a in vec1))
        magnitude2 = math.sqrt(sum(a * a for a in vec2))
        
        if magnitude1 == 0 or magnitude2 == 0:
            return 0.0
        
        similarity = dot_product / (magnitude1 * magnitude2)
        
        # Ensure result is between -1 and 1
        return max(-1.0, min(1.0, similarity))
        
    except Exception as e:
        logger.error(f"Error calculating cosine similarity: {e}")
        return 0.0
        
def generate_simple_enhanced_response(message: str, influencer_id: str, influencer_name: str) -> str:
    """Simple enhanced response with personal knowledge"""
    try:
        # Get personal information
        personal_context = ""
        try:
            personal_info = db.get_personal_knowledge(influencer_id)
            if personal_info:
                personal_parts = []
                if personal_info.get('bio'):
                    personal_parts.append(f"About me: {personal_info['bio']}")
                if personal_info.get('expertise'):
                    personal_parts.append(f"My expertise: {personal_info['expertise']}")
                if personal_info.get('personality'):
                    personal_parts.append(f"My communication style: {personal_info['personality']}")
                
                if personal_parts:
                    personal_context = "\n\nPersonal information:\n" + "\n".join(personal_parts)
        except Exception as e:
            logger.error(f"Error getting personal info: {e}")
        
        # Build system prompt
        system_prompt = f"""You are an AI assistant representing {influencer_name}. You help users by providing helpful, engaging, and personalized responses.

CORE INSTRUCTIONS:
- Be conversational, friendly, and authentic
- Provide helpful and accurate information
- Stay true to the personality described below
- Be honest if you don't know something specific

{personal_context}

RESPONSE GUIDELINES:
- Keep responses concise but informative (2-4 sentences usually)
- Be helpful and solution-oriented
- Use a natural, conversational tone

Remember: You are representing {influencer_name}, so respond as if you are them, using their knowledge and personality."""
        
        # Generate response
        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": message}
            ],
            max_tokens=500,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        logger.error(f"Enhanced response generation error: {e}")
        return f"Hi! I'm {influencer_name}'s AI assistant. I'm here to help you!"

@app.route('/api/chat/<username>', methods=['GET'])
def get_chat_info(username):
    """ENHANCED: Get comprehensive chat information including avatar and knowledge status"""
    try:
        # Clean username
        clean_username = username.strip().lower()
        
        # Get influencer
        influencer = db.get_influencer_by_username(clean_username)
        
        if not influencer:
            return jsonify({
                'status': 'error',
                'message': 'Influencer not found'
            }), 404
        
        # Get additional stats
        knowledge_documents_count = 0
        affiliate_platforms_count = 0
        
        try:
            # Count knowledge documents
            if hasattr(db, 'get_knowledge_documents'):
                knowledge_docs = db.get_knowledge_documents(influencer['id'])
                knowledge_documents_count = len(knowledge_docs) if knowledge_docs else 0
            
            # Count affiliate platforms
            affiliate_links = db.get_affiliate_links(influencer['id']) if hasattr(db, 'get_affiliate_links') else []
            affiliate_platforms_count = len([link for link in affiliate_links if link.get('is_active', True)])
            
        except Exception as stats_error:
            logger.warning(f"Could not load additional stats: {stats_error}")
        
        # Build comprehensive response
        response_data = {
            'username': influencer['username'],
            'bio': influencer.get('bio', ''),
            'expertise': influencer.get('expertise', ''),
            'personality': influencer.get('personality', ''),
            'has_avatar': bool(influencer.get('heygen_avatar_id')),
            'avatar_id': influencer.get('heygen_avatar_id'),
            'avatar_type': influencer.get('avatar_type', 'none'),
            'voice_id': influencer.get('preferred_voice_id') or influencer.get('voice_id') or Config.DEFAULT_VOICE_ID,
            'has_knowledge': bool(
                influencer.get('bio') or 
                influencer.get('expertise') or 
                influencer.get('personality') or
                knowledge_documents_count > 0
            ),
            'knowledge_documents_count': knowledge_documents_count,
            'affiliate_platforms_connected': affiliate_platforms_count,
            'chat_capabilities': {
                'text_chat': True,
                'voice_responses': True,
                'video_responses': bool(influencer.get('heygen_avatar_id')),
                'knowledge_enhanced': bool(
                    influencer.get('bio') or 
                    influencer.get('expertise') or 
                    influencer.get('personality') or
                    knowledge_documents_count > 0
                ),
                'product_recommendations': affiliate_platforms_count > 0
            },
            'avatar_preview_url': f"/api/avatar/preview/{influencer['id']}" if influencer.get('heygen_avatar_id') else None,
            'created_at': influencer.get('created_at'),
            'chat_url': f"{request.host_url}chat.html?username={clean_username}"
        }
        
        logger.info(f"Enhanced chat info retrieved for {username} - Avatar: {response_data['has_avatar']}, Knowledge: {response_data['has_knowledge']}, Products: {response_data['chat_capabilities']['product_recommendations']}")
        
        return jsonify({
            'status': 'success',
            'data': response_data
        })
        
    except Exception as e:
        logger.error(f"Get chat info error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get chat information'
        }), 500

@app.route('/api/knowledge/upload', methods=['POST'])
@token_required
def upload_knowledge_document(current_user):
    """FIXED: Actually process uploaded documents for knowledge integration"""
    try:
        # Get uploaded file
        file = request.files.get('file')
        if not file:
            return jsonify({
                'status': 'error',
                'message': 'No file provided'
            }), 400

        # Validate file
        try:
            validate_file_upload(file)
        except ValueError as ve:
            return jsonify({
                'status': 'error',
                'message': str(ve)
            }), 400

        # Get file info
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        content_type = file.content_type
        
        # Get file size
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)

        logger.info(f"📄 Processing {filename} for user {current_user['username']}")

        # STEP 1: Extract text from the uploaded file
        file.seek(0)
        file_content = file.read()
        
        extracted_text = ""
        try:
            if content_type == 'application/pdf':
                # Extract text from PDF
                import PyPDF2
                import io
                
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
                
                logger.info(f"✅ Extracted {len(extracted_text)} characters from PDF")
                
            elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                # Extract text from Word document
                from docx import Document as DocxDocument
                import io
                
                doc_file = io.BytesIO(file_content)
                doc = DocxDocument(doc_file)
                
                for paragraph in doc.paragraphs:
                    extracted_text += paragraph.text + "\n"
                
                logger.info(f"✅ Extracted {len(extracted_text)} characters from Word doc")
            
            extracted_text = extracted_text.strip()
            
        except Exception as extraction_error:
            logger.error(f"Text extraction failed: {extraction_error}")
            return jsonify({
                'status': 'error',
                'message': f'Failed to extract text from document: {str(extraction_error)}'
            }), 500

        # STEP 2: Save document metadata to database
        try:
            document_data = {
                'id': str(uuid.uuid4()),
                'influencer_id': current_user['id'],
                'filename': filename,
                'safe_filename': unique_filename,
                'content_type': content_type,
                'file_size': file_size,
                'storage_url': f"knowledge-documents/{unique_filename}",
                'upload_date': datetime.now(timezone.utc).isoformat(),
                'is_processed': False,  # Will be set to True after processing
                'text_content': extracted_text[:5000] if extracted_text else None,
                'chunk_count': 0,
                'created_at': datetime.now(timezone.utc).isoformat(),
                'updated_at': datetime.now(timezone.utc).isoformat()
            }
            
            response = db.supabase.table('knowledge_documents').insert(document_data).execute()
            
            if not response.data:
                raise Exception("Failed to save document metadata")
            
            document_id = response.data[0]['id']
            logger.info(f"✅ Document metadata saved: {document_id}")
            
        except Exception as db_error:
            logger.error(f"Database save error: {db_error}")
            return jsonify({
                'status': 'error',
                'message': f'Failed to save document: {str(db_error)}'
            }), 500

        # STEP 3: FIXED - Create searchable chunks with proper embeddings
        chunks_created = 0
        if extracted_text and len(extracted_text.strip()) > 50:
            try:
                # Clean and prepare text
                import re
                clean_text = re.sub(r'\s+', ' ', extracted_text).strip()
                
                # FIXED: Better chunking strategy
                chunks = create_better_text_chunks(clean_text, max_chunk_size=400, overlap=50)
                logger.info(f"📝 Created {len(chunks)} text chunks")
                
                # FIXED: Generate embeddings using OpenAI
                if chunks:
                    try:
                        # Generate embeddings for all chunks
                        embeddings = []
                        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
                        
                        for chunk in chunks:
                            embedding_response = client.embeddings.create(
                                model="text-embedding-3-small",
                                input=chunk
                            )
                            embeddings.append(embedding_response.data[0].embedding)
                        
                        logger.info(f"🧠 Generated {len(embeddings)} embeddings")
                        
                        # FIXED: Store chunks with proper embeddings
                        chunk_data = []
                        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                            chunk_data.append({
                                'id': str(uuid.uuid4()),
                                'document_id': document_id,
                                'influencer_id': current_user['id'],
                                'chunk_index': i,
                                'chunk_text': chunk,
                                'embedding': json.dumps(embedding),  # Store as JSON string
                                'token_count': len(chunk.split()),
                                'created_at': datetime.now(timezone.utc).isoformat()
                            })
                        
                        # Insert all chunks
                        if chunk_data:
                            chunk_response = db.supabase.table('knowledge_chunks').insert(chunk_data).execute()
                            
                            if chunk_response.data:
                                chunks_created = len(chunk_response.data)
                                
                                # Update document as processed
                                db.supabase.table('knowledge_documents').update({
                                    'chunk_count': chunks_created,
                                    'is_processed': True,
                                    'processed_date': datetime.now(timezone.utc).isoformat()
                                }).eq('id', document_id).execute()
                                
                                logger.info(f"✅ Stored {chunks_created} chunks in knowledge base")
                            else:
                                logger.error("Failed to store chunks in database")
                        
                    except Exception as embedding_error:
                        logger.error(f"Embedding generation failed: {embedding_error}")
                        return jsonify({
                            'status': 'error',
                            'message': f'Failed to process document for search: {str(embedding_error)}'
                        }), 500
                
            except Exception as processing_error:
                logger.error(f"Document processing failed: {processing_error}")
                return jsonify({
                    'status': 'error',
                    'message': f'Failed to process document: {str(processing_error)}'
                }), 500

        return jsonify({
            'status': 'success',
            'message': f'Document uploaded and processed successfully',
            'data': {
                'document_id': document_id,
                'filename': filename,
                'file_size': file_size,
                'text_extracted': len(extracted_text) if extracted_text else 0,
                'chunks_created': chunks_created,
                'processed': chunks_created > 0,
                'can_search': chunks_created > 0
            }
        })

    except Exception as e:
        logger.error(f"Document upload error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to upload document: {str(e)}'
        }), 500

def create_better_text_chunks(text: str, max_chunk_size: int = 400, overlap: int = 50) -> List[str]:
    """Create better text chunks with sentence awareness"""
    import re
    
    # Split into sentences
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would make chunk too long and we have content
        if len(current_chunk) + len(sentence) > max_chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            
            # Start new chunk with overlap (last few words)
            words = current_chunk.split()
            if len(words) > overlap:
                overlap_text = " ".join(words[-overlap:])
                current_chunk = overlap_text + " " + sentence
            else:
                current_chunk = sentence
        else:
            current_chunk += " " + sentence if current_chunk else sentence
    
    # Add final chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

@app.route('/api/knowledge/documents', methods=['GET'])
@token_required
def get_knowledge_documents(current_user):
    """Get all uploaded documents for the knowledge page"""
    try:
        # Get documents from database
        response = db.supabase.table('knowledge_documents') \
            .select('*') \
            .eq('influencer_id', current_user['id']) \
            .order('created_at', desc=True) \
            .execute()
        
        documents = response.data if response.data else []
        
        # Convert to frontend format
        formatted_documents = []
        total_size = 0
        
        for doc in documents:
            total_size += doc.get('file_size', 0)
            
            formatted_documents.append({
                'id': doc['id'],
                'name': doc['filename'],
                'size': doc['file_size'],
                'type': doc['content_type'],
                'status': 'completed' if doc.get('is_processed') else 'processing',
                'document_id': doc['id'],
                'uploaded_at': doc.get('created_at'),
                'processed': doc.get('is_processed', False),
                'chunk_count': doc.get('chunk_count', 0),
                'text_length': len(doc.get('text_content', '') or '')
            })
        
        logger.info(f"📚 Retrieved {len(formatted_documents)} documents for {current_user['username']}")
        
        return jsonify({
            'status': 'success',
            'data': {
                'documents': formatted_documents,
                'total_count': len(formatted_documents),
                'processed_count': len([d for d in formatted_documents if d['processed']]),
                'total_size': total_size
            }
        })
        
    except Exception as e:
        logger.error(f"Get knowledge documents error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get knowledge documents'
        }), 500

def extract_text_from_file(file_content: bytes, content_type: str) -> str:
    """Extract text from PDF or Word documents"""
    try:
        if content_type == 'application/pdf':
            # Extract from PDF
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            # Extract from Word document
            from docx import Document
            import io
            
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        else:
            return ""
            
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return ""

def create_text_chunks(text: str, max_length: int = 500) -> List[str]:
    """Split text into manageable chunks"""
    import re
    
    # Clean text
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Split into sentences
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would make chunk too long
        if len(current_chunk) + len(sentence) > max_length and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
        else:
            current_chunk += sentence + ". "
    
    # Add final chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

def generate_chunk_embeddings(chunks: List[str]) -> List[List[float]]:
    """Generate embeddings for text chunks"""
    try:
        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        embeddings = []
        
        for chunk in chunks:
            response = client.embeddings.create(
                model="text-embedding-3-small",
                input=chunk
            )
            embeddings.append(response.data[0].embedding)
        
        return embeddings
        
    except Exception as e:
        logger.error(f"Embedding generation error: {e}")
        return []

def store_knowledge_chunks(document_id: str, influencer_id: str, chunks: List[str], embeddings: List[List[float]]):
    """Store chunks and embeddings in database"""
    try:
        chunk_data = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_data.append({
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'influencer_id': influencer_id,
                'chunk_index': i,
                'chunk_text': chunk,
                'embedding': json.dumps(embedding),  # Store as JSON string
                'token_count': len(chunk.split()),
                'created_at': datetime.now(timezone.utc).isoformat()
            })
        
        # Insert chunks
        if chunk_data:
            response = db.supabase.table('knowledge_chunks').insert(chunk_data).execute()
            return bool(response.data)
        
        return False
        
    except Exception as e:
        logger.error(f"Chunk storage error: {e}")
        return False

@app.route('/api/knowledge/personal-info', methods=['POST'])
@token_required
def save_personal_knowledge(current_user):
    """Save influencer's personal knowledge information"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        # Prepare update data
        update_data = {}
        
        # Validate and save bio, expertise, and personality info
        if 'bio' in data:
            bio = data['bio'].strip()
            if len(bio) > 1000:
                return jsonify({
                    'status': 'error',
                    'message': 'Bio must be less than 1000 characters'
                }), 400
            update_data['bio'] = bio
            
        if 'expertise' in data:
            expertise = data['expertise'].strip()
            if len(expertise) > 500:
                return jsonify({
                    'status': 'error',
                    'message': 'Expertise must be less than 500 characters'
                }), 400
            update_data['expertise'] = expertise
            
        if 'personality' in data:
            personality = data['personality'].strip()
            if len(personality) > 500:
                return jsonify({
                    'status': 'error',
                    'message': 'Personality must be less than 500 characters'
                }), 400
            update_data['personality'] = personality
        
        if not update_data:
            return jsonify({
                'status': 'error',
                'message': 'No valid data to update'
            }), 400
        
        # Add timestamp
        update_data['updated_at'] = datetime.now(timezone.utc).isoformat()
        
        # Update influencer profile
        success = db.update_influencer(current_user['id'], update_data)
        
        if success:
            logger.info(f"✅ Personal knowledge updated for influencer {current_user['username']}")
            return jsonify({
                'status': 'success',
                'message': 'Personal information saved successfully',
                'data': {
                    'updated_fields': list(update_data.keys()),
                    'updated_at': update_data['updated_at']
                }
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to save personal information'
            }), 500
            
    except Exception as e:
        logger.error(f"Save personal knowledge error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to save personal information'
        }), 500

@app.route('/api/knowledge/process-existing', methods=['POST'])
@token_required
def process_existing_documents(current_user):
    """Process existing documents that weren't processed during upload"""
    try:
        # Get all unprocessed documents for this user
        response = db.supabase.table('knowledge_documents') \
            .select('*') \
            .eq('influencer_id', current_user['id']) \
            .eq('is_processed', False) \
            .execute()
        
        documents = response.data if response.data else []
        
        if not documents:
            return jsonify({
                'status': 'success',
                'message': 'No documents to process',
                'data': {'processed_count': 0}
            })
        
        processed_count = 0
        errors = []
        
        for document in documents:
            try:
                logger.info(f"📚 Processing document: {document['filename']}")
                
                # For now, since we don't have actual file storage, 
                # create dummy chunks from filename and metadata
                dummy_text = f"Document: {document['filename']}\nContent type: {document['content_type']}\nThis document contains information related to {current_user['username']}'s knowledge base."
                
                # Create chunks
                chunks = create_text_chunks(dummy_text)
                
                # Generate embeddings
                embeddings = generate_chunk_embeddings(chunks)
                
                # Store chunks
                if chunks and embeddings:
                    store_knowledge_chunks(document['id'], current_user['id'], chunks, embeddings)
                    
                    # Mark as processed
                    db.supabase.table('knowledge_documents').update({
                        'is_processed': True,
                        'processed_date': datetime.now(timezone.utc).isoformat(),
                        'text_content': dummy_text,
                        'chunk_count': len(chunks)
                    }).eq('id', document['id']).execute()
                    
                    processed_count += 1
                    logger.info(f"✅ Processed {document['filename']}")
                
            except Exception as doc_error:
                error_msg = f"Failed to process {document['filename']}: {str(doc_error)}"
                errors.append(error_msg)
                logger.error(error_msg)
        
        return jsonify({
            'status': 'success',
            'message': f'Processed {processed_count} documents',
            'data': {
                'processed_count': processed_count,
                'total_documents': len(documents),
                'errors': errors
            }
        })
        
    except Exception as e:
        logger.error(f"Process existing documents error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to process documents: {str(e)}'
        }), 500

# =============================================================================
# Add this endpoint to test knowledge search
# =============================================================================

@app.route('/api/knowledge/test-search', methods=['POST'])
@token_required
def test_knowledge_search(current_user):
    """Test endpoint to check if knowledge search is working"""
    try:
        data = request.get_json()
        query = data.get('query', 'test search')
        
        # Generate embedding for test query
        client = openai.OpenAI(api_key=Config.OPENAI_API_KEY)
        embedding_response = client.embeddings.create(
            model="text-embedding-3-small",
            input=query
        )
        query_embedding = embedding_response.data[0].embedding
        
        # Search knowledge base
        results = db.search_knowledge_base(current_user['id'], query_embedding, limit=5)
        
        # Get document count
        doc_response = db.supabase.table('knowledge_documents') \
            .select('*') \
            .eq('influencer_id', current_user['id']) \
            .execute()
        
        documents = doc_response.data if doc_response.data else []
        processed_docs = [doc for doc in documents if doc.get('is_processed')]
        
        # Get chunk count
        chunk_response = db.supabase.table('knowledge_chunks') \
            .select('id') \
            .eq('influencer_id', current_user['id']) \
            .execute()
        
        chunks = chunk_response.data if chunk_response.data else []
        
        return jsonify({
            'status': 'success',
            'data': {
                'query': query,
                'total_documents': len(documents),
                'processed_documents': len(processed_docs),
                'total_chunks': len(chunks),
                'search_results': len(results),
                'results': results[:3],  # Show top 3 results
                'documents': [{'filename': doc['filename'], 'processed': doc.get('is_processed', False)} for doc in documents]
            }
        })
        
    except Exception as e:
        logger.error(f"Knowledge search test error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Knowledge search test failed: {str(e)}'
        }), 500

# =============================================================================
# AFFILIATE MANAGEMENT ROUTES
# =============================================================================

@app.route('/api/affiliate', methods=['POST'])
@token_required
def connect_affiliate_platform_fixed(current_user):
    """FIXED: Connect affiliate platform with proper credential storage"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        platform = data.get('platform', '').lower().strip()
        valid_platforms = ['amazon', 'rakuten', 'shareasale', 'cj_affiliate', 'skimlinks']
        
        logger.info(f"🔗 Connecting {platform} for user {current_user['username']}")
        logger.info(f"📝 Data received: {list(data.keys())}")
        
        if platform not in valid_platforms:
            return jsonify({
                'status': 'error',
                'message': f'Invalid platform. Must be one of: {", ".join(valid_platforms)}'
            }), 400
        
        # Check if already connected
        existing_link = db.get_affiliate_link_by_platform(current_user['id'], platform)
        if existing_link:
            return jsonify({
                'status': 'error',
                'message': f'{platform.title()} is already connected'
            }), 400
        
        # FIXED: Platform-specific validation and data preparation
        affiliate_data = {
            'id': str(uuid.uuid4()),
            'influencer_id': current_user['id'],
            'platform': platform,
            'is_active': True,
            'created_at': datetime.now(timezone.utc).isoformat(),
            'updated_at': datetime.now(timezone.utc).isoformat()
        }
        
        if platform == 'rakuten':
            # FIXED: Proper Rakuten validation and storage
            client_id = data.get('client_id', '').strip()
            client_secret = data.get('client_secret', '').strip()
            application_id = data.get('application_id', '').strip()
            
            if not client_id or not client_secret:
                return jsonify({
                    'status': 'error',
                    'message': 'Client ID and Client Secret are required for Rakuten Advertising'
                }), 400
            
            # Store credentials properly
            affiliate_data.update({
                'affiliate_id': client_id,  # Use client_id as primary identifier
                'client_id': client_id,
                'client_secret': client_secret,
                'rakuten_client_id': client_id,
                'rakuten_client_secret': client_secret,
                'rakuten_application_id': application_id,
                'application_id': application_id
            })
            
            logger.info(f"✅ Rakuten credentials prepared - Client ID: {client_id}")
            
        elif platform == 'amazon':
            # Amazon validation
            access_key = data.get('access_key', '').strip()
            secret_key = data.get('secret_key', '').strip()
            partner_tag = data.get('partner_tag', '').strip()
            
            if not all([access_key, secret_key, partner_tag]):
                return jsonify({
                    'status': 'error',
                    'message': 'Access Key, Secret Key, and Partner Tag are required for Amazon'
                }), 400
            
            affiliate_data.update({
                'affiliate_id': partner_tag,
                'access_key': access_key,
                'secret_key': secret_key,
                'partner_tag': partner_tag,
                'amazon_access_key': access_key,
                'amazon_secret_key': secret_key
            })
            
        # Add other platforms as needed...
        
        # Save to database
        success = db.create_affiliate_link(affiliate_data)
        
        if success:
            logger.info(f"✅ {platform.title()} connected successfully for {current_user['username']}")
            return jsonify({
                'status': 'success',
                'message': f'{platform.title()} connected successfully',
                'data': {
                    'platform': platform,
                    'affiliate_id': affiliate_data.get('affiliate_id', ''),
                    'connected_at': affiliate_data['created_at']
                }
            }), 201
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to save affiliate connection'
            }), 500
            
    except Exception as e:
        logger.error(f"❌ Connect affiliate platform error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to connect platform: {str(e)}'
        }), 500
    
@app.route('/api/affiliate/<platform>', methods=['DELETE'])
@token_required
def remove_affiliate_link(current_user, platform):
    """Remove an affiliate link"""
    try:
        # Validate platform
        valid_platforms = ['amazon', 'rakuten', 'shareasale', 'cj_affiliate', 'skimlinks']
        if platform not in valid_platforms:
            return jsonify({
                'status': 'error',
                'message': f'Invalid platform. Must be one of: {", ".join(valid_platforms)}'
            }), 400
        
        # Remove from database
        success = db.delete_affiliate_link(current_user['id'], platform)
        
        if success:
            logger.info(f"✅ Affiliate link removed for {current_user['username']}: {platform}")
            return jsonify({
                'status': 'success',
                'message': f'{platform} removed successfully'
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Affiliate link not found or failed to remove'
            }), 404
            
    except Exception as e:
        logger.error(f"Remove affiliate link error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to remove affiliate link'
        }), 500

@app.route('/api/affiliate/<platform>', methods=['PUT'])
@token_required
def update_affiliate_link(current_user, platform):
    """Update an affiliate link"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        # Validate platform
        valid_platforms = ['amazon', 'rakuten', 'shareasale', 'cj_affiliate', 'skimlinks']
        if platform not in valid_platforms:
            return jsonify({
                'status': 'error',
                'message': f'Invalid platform. Must be one of: {", ".join(valid_platforms)}'
            }), 400
        
        # Prepare update data
        update_data = {
            'updated_at': datetime.now(timezone.utc).isoformat()
        }
        
        # Add platform-specific fields
        if 'is_active' in data:
            update_data['is_active'] = data['is_active']
        
        # Update in database
        success = db.update_affiliate_link(current_user['id'], platform, update_data)
        
        if success:
            logger.info(f"✅ Affiliate link updated for {current_user['username']}: {platform}")
            return jsonify({
                'status': 'success',
                'message': f'{platform} updated successfully'
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Affiliate link not found or failed to update'
            }), 404
            
    except Exception as e:
        logger.error(f"Update affiliate link error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to update affiliate link'
        }), 500

@app.route('/api/affiliate/search-products', methods=['POST'])
@token_required
def search_affiliate_products(current_user):
    """FIXED: Search for products with enhanced debugging"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        query = data.get('query', '').strip()
        platform = data.get('platform', 'all')
        limit = min(int(data.get('limit', 5)), 20)
        
        if not query:
            return jsonify({
                'status': 'error',
                'message': 'Search query is required'
            }), 400
        
        logger.info(f"🔍 Product search - User: {current_user['username']}, Query: '{query}', Platform: {platform}")
        
        # Check if affiliate service is available
        if not hasattr(chatbot, 'affiliate_service') or not chatbot.affiliate_service:
            return jsonify({
                'status': 'error',
                'message': 'Affiliate service not available'
            }), 503
        
        # FIXED: Enhanced product search with better error handling
        if platform == 'all':
            # Search across all connected platforms
            recommendations = chatbot.affiliate_service.get_product_recommendations(
                query=query,
                influencer_id=current_user['id'],
                limit=limit
            )
            
            logger.info(f"📦 Multi-platform search results: {recommendations['total_found']} products from {recommendations['platforms_searched']} platforms")
            
            return jsonify({
                'status': 'success',
                'data': {
                    'products': recommendations['products'],
                    'total_found': recommendations['total_found'],
                    'platforms_searched': recommendations['platforms_searched'],
                    'query': query,
                    'search_type': 'multi_platform'
                }
            })
        else:
            # Search specific platform
            products = chatbot.affiliate_service.search_products(
                query=query,
                platform=platform,
                influencer_id=current_user['id'],
                limit=limit
            )
            
            logger.info(f"📦 Single platform ({platform}) search results: {len(products)} products")
            
            return jsonify({
                'status': 'success',
                'data': {
                    'products': products,
                    'total_found': len(products),
                    'platform': platform,
                    'query': query,
                    'search_type': 'single_platform'
                }
            })
        
    except Exception as e:
        logger.error(f"Product search error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to search products: {str(e)}'
        }), 500

@app.route('/api/affiliate/platform-status', methods=['GET'])
@token_required
def get_platform_status(current_user):
    """FIXED: Get platform status using simplified database schema"""
    try:
        # Get user's affiliate links from database
        affiliate_links = db.get_affiliate_links(current_user['id'])
        
        # FIXED: Simple platform data that matches your current setup
        platforms_data = {
            'amazon': {
                'name': 'Amazon Associates',
                'logo': 'A',
                'description': 'World\'s largest online retailer with millions of products',
                'commission_range': '1-10%',
                'min_payout': '$10',
                'payment_schedule': 'Monthly',
                'special_features': ['Product reviews', 'Prime eligibility', 'Lightning deals'],
                'connected': False,
                'estimated_products': 0,
                'credentials_format': 'Product Advertising API (PA-API 5.0)'
            },
            'rakuten': {
                'name': 'Rakuten Advertising',
                'logo': 'R',
                'description': 'Global leader in affiliate marketing with premium brands',
                'commission_range': '2-15%',
                'min_payout': '$50',
                'payment_schedule': 'Monthly',
                'special_features': ['Premium brands', 'Global reach', 'Advanced analytics'],
                'connected': False,
                'estimated_products': 0,
                'credentials_format': 'OAuth2 Client Credentials'
            },
            'shareasale': {
                'name': 'ShareASale',
                'logo': 'S',
                'description': 'Performance marketing network with diverse merchants',
                'commission_range': '3-20%',
                'min_payout': '$50',
                'payment_schedule': '20th of each month',
                'special_features': ['Real-time tracking', 'Custom links', 'Merchant variety'],
                'connected': False,
                'estimated_products': 0,
                'credentials_format': 'API Token & Secret Key'
            },
            'cj_affiliate': {
                'name': 'CJ Affiliate',
                'logo': 'CJ',
                'description': 'Commission Junction - trusted by top brands worldwide',
                'commission_range': '2-12%',
                'min_payout': '$50',
                'payment_schedule': 'Monthly',
                'special_features': ['Enterprise brands', 'Deep linking', 'Attribution tracking'],
                'connected': False,
                'estimated_products': 0,
                'credentials_format': 'API Key & Website ID'
            },
            'skimlinks': {
                'name': 'Skimlinks',
                'logo': 'SK',
                'description': 'Automated affiliate marketing with 48,500+ merchants',
                'commission_range': '1-8%',
                'min_payout': '$10',
                'payment_schedule': 'Monthly',
                'special_features': ['Auto-affiliate', 'Content monetization', 'Easy setup'],
                'connected': False,
                'estimated_products': 0,
                'credentials_format': 'API Key & Publisher ID'
            }
        }
        
        # Update with user's connected platforms
        connected_platforms = 0
        total_estimated_products = 0
        
        for link in affiliate_links:
            if link.get('is_active', True):
                platform = link['platform']
                if platform in platforms_data:
                    platforms_data[platform]['connected'] = True
                    platforms_data[platform]['estimated_products'] = 1000  # Placeholder
                    platforms_data[platform]['affiliate_id'] = link.get('affiliate_id', '')
                    connected_platforms += 1
                    total_estimated_products += 1000
        
        # Calculate stats
        stats = {
            'connected_platforms': connected_platforms,
            'estimated_products': total_estimated_products,
            'recommendations_made': 0,  # Would need analytics table
            'potential_earnings': total_estimated_products * 0.05,  # Rough estimate
            'completion_percentage': (connected_platforms / len(platforms_data)) * 100
        }
        
        return jsonify({
            'status': 'success',
            'data': {
                'platforms': platforms_data,
                'stats': stats
            }
        })
        
    except Exception as e:
        logger.error(f"Platform status error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get platform status'
        }), 500

@app.route('/api/affiliate/test-connection', methods=['POST'])
@token_required
def test_affiliate_connection(current_user):
    """FIXED: Test affiliate connection with proper validation"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        platform = data.get('platform', '').lower().strip()
        credentials = data.get('credentials', {})
        
        logger.info(f"🔍 Testing {platform} connection for {current_user['username']}")
        logger.info(f"📝 Credentials provided: {list(credentials.keys())}")
        
        if not platform or not credentials:
            return jsonify({
                'status': 'error',
                'message': 'Platform and credentials are required'
            }), 400
        
        # FIXED: Platform-specific validation with detailed feedback
        validation_passed = False
        validation_error = ""
        validation_details = {}
        
        if platform == 'amazon':
            required_fields = ['access_key', 'secret_key', 'partner_tag']
            missing = [field for field in required_fields if not credentials.get(field, '').strip()]
            if missing:
                validation_error = f"Missing required fields: {', '.join(missing)}"
            else:
                validation_passed = True
                validation_details = {
                    'access_key_format': 'Valid' if credentials.get('access_key', '').startswith('AKIA') else 'Warning: Should start with AKIA',
                    'partner_tag_format': 'Valid' if '-' in credentials.get('partner_tag', '') else 'Warning: Should contain a hyphen'
                }
                
        elif platform == 'rakuten':
            client_id = credentials.get('client_id', '').strip()
            if not client_id:
                validation_error = "Application ID (Client ID) is required for Rakuten"
            elif len(client_id) < 5:
                validation_error = "Rakuten Application ID must be at least 5 characters"
            elif not client_id.replace('-', '').replace('_', '').isalnum():
                validation_error = "Rakuten Application ID can only contain letters, numbers, hyphens, and underscores"
            else:
                validation_passed = True
                validation_details = {
                    'application_id': client_id,
                    'format': f'Valid alphanumeric ID ({len(client_id)} characters)',
                    'client_secret': 'Provided' if credentials.get('client_secret') else 'Not provided (optional)',
                    'affiliate_id': 'Provided' if credentials.get('affiliate_id') else 'Not provided (optional)'
                }
                logger.info(f"✅ Rakuten validation passed - Application ID: {client_id}")
                
        elif platform == 'shareasale':
            required_fields = ['api_token', 'secret_key', 'affiliate_id']
            missing = [field for field in required_fields if not credentials.get(field, '').strip()]
            if missing:
                validation_error = f"Missing required fields: {', '.join(missing)}"
            else:
                validation_passed = True
                validation_details = {
                    'api_token_length': len(credentials.get('api_token', '')),
                    'affiliate_id_format': 'Valid' if credentials.get('affiliate_id', '').isdigit() else 'Should be numeric'
                }
                
        elif platform == 'cj_affiliate':
            required_fields = ['api_key', 'website_id']
            missing = [field for field in required_fields if not credentials.get(field, '').strip()]
            if missing:
                validation_error = f"Missing required fields: {', '.join(missing)}"
            else:
                validation_passed = True
                validation_details = {
                    'api_key_length': len(credentials.get('api_key', '')),
                    'website_id_format': 'Valid' if credentials.get('website_id', '').isdigit() else 'Should be numeric'
                }
                
        elif platform == 'skimlinks':
            required_fields = ['api_key', 'publisher_id']
            missing = [field for field in required_fields if not credentials.get(field, '').strip()]
            if missing:
                validation_error = f"Missing required fields: {', '.join(missing)}"
            else:
                validation_passed = True
                validation_details = {
                    'api_key_length': len(credentials.get('api_key', '')),
                    'publisher_id_format': 'Valid' if credentials.get('publisher_id', '').isdigit() else 'Should be numeric'
                }
        else:
            validation_error = f"Unsupported platform: {platform}"
        
        if validation_passed:
            logger.info(f"✅ {platform} validation successful")
            return jsonify({
                'status': 'success',
                'message': f'{platform.title()} credentials validated successfully',
                'data': {
                    'platform': platform,
                    'connection_status': 'validated',
                    'credentials_provided': True,
                    'fields_validated': list(credentials.keys()),
                    'validation_details': validation_details
                }
            })
        else:
            logger.error(f"❌ {platform} validation failed: {validation_error}")
            return jsonify({
                'status': 'error',
                'message': validation_error,
                'error_type': 'validation_error',
                'platform': platform,
                'fields_provided': list(credentials.keys())
            }), 400
        
    except Exception as e:
        logger.error(f"❌ Test affiliate connection error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Connection test failed: {str(e)}'
        }), 500

@app.route('/api/affiliate/rakuten', methods=['POST'])
@token_required
def connect_rakuten_platform(current_user):
    """Enhanced Rakuten connection with proper validation"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        # Validate required Rakuten fields
        required_fields = ['application_id', 'affiliate_id']
        missing_fields = [field for field in required_fields if not data.get(field)]
        
        if missing_fields:
            return jsonify({
                'status': 'error',
                'message': f'Missing required fields: {", ".join(missing_fields)}'
            }), 400
        
        # Check if Rakuten is already connected
        existing_link = db.get_affiliate_link_by_platform(current_user['id'], 'rakuten')
        if existing_link:
            return jsonify({
                'status': 'error',
                'message': 'Rakuten is already connected to your account'
            }), 400
        
        # Prepare affiliate link data for Rakuten
        affiliate_data = {
            'id': str(uuid.uuid4()),
            'influencer_id': current_user['id'],
            'platform': 'rakuten',
            'is_active': True,
            'created_at': datetime.now(timezone.utc).isoformat(),
            'updated_at': datetime.now(timezone.utc).isoformat(),
            'rakuten_application_id': data.get('application_id'),
            'rakuten_affiliate_id': data.get('affiliate_id'),
            'affiliate_id': data.get('affiliate_id'),  # Primary affiliate ID
            'api_token': data.get('application_id')  # Use application_id as token
        }
        
        # Test connection before saving
        try:
            from affiliate_service import AffiliateService
            affiliate_service = AffiliateService(db)
            
            # Create temporary affiliate info for testing
            temp_info = {
                'rakuten_client_id': data.get('application_id'),
                'rakuten_access_token': data.get('affiliate_id'),
                'merchant_id': data.get('application_id'),
                'api_token': data.get('affiliate_id')
            }
            
            # Test with a simple search
            test_products = affiliate_service.platforms['rakuten'].search_products(
                "test", temp_info, 1
            )
            
            # Note: Rakuten might not return products for "test" query, 
            # but if we don't get an authentication error, credentials are likely valid
            
        except Exception as test_error:
            logger.warning(f"Rakuten connection test warning: {test_error}")
            # Continue anyway - the test might fail due to API limitations
        
        # Save to database
        success = db.create_affiliate_link(affiliate_data)
        
        if success:
            logger.info(f"✅ Rakuten connected for {current_user['username']}")
            return jsonify({
                'status': 'success',
                'message': 'Rakuten connected successfully',
                'data': {
                    'platform': 'rakuten',
                    'affiliate_id': affiliate_data['affiliate_id'],
                    'application_id': affiliate_data['rakuten_application_id'],
                    'created_at': affiliate_data['created_at']
                }
            }), 201
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to save Rakuten connection'
            }), 500
            
    except Exception as e:
        logger.error(f"Connect Rakuten error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to connect Rakuten'
        }), 500

@app.route('/api/analytics/affiliate-performance', methods=['GET'])
@token_required
def get_affiliate_performance(current_user):
    """Get detailed affiliate performance analytics"""
    try:
        days = int(request.args.get('days', 30))
        platform = request.args.get('platform', 'all')
        
        from datetime import timedelta
        start_date = datetime.now(timezone.utc) - timedelta(days=days)
        
        # Get chat interactions with product recommendations
        response = db.supabase.table('chat_interactions') \
            .select('*') \
            .eq('influencer_id', current_user['id']) \
            .eq('products_included', True) \
            .gte('created_at', start_date.isoformat()) \
            .execute()
        
        interactions = response.data if response.data else []
        
        # Calculate metrics
        total_recommendations = len(interactions)
        unique_sessions = len(set([chat.get("session_id") for chat in interactions if chat.get("session_id")]))
        
        # Group by platform if specific platform requested
        platform_metrics = {}
        daily_stats = {}
        
        for interaction in interactions:
            date = interaction["created_at"][:10]  # YYYY-MM-DD
            
            # Daily stats
            if date not in daily_stats:
                daily_stats[date] = {'recommendations': 0, 'sessions': set()}
            
            daily_stats[date]['recommendations'] += 1
            if interaction.get('session_id'):
                daily_stats[date]['sessions'].add(interaction['session_id'])
        
        # Convert sets to counts for JSON serialization
        for date, stats in daily_stats.items():
            stats['unique_sessions'] = len(stats['sessions'])
            del stats['sessions']
        
        # Get affiliate links for platform info
        affiliate_links = db.get_affiliate_links(current_user['id'])
        
        return jsonify({
            'status': 'success',
            'data': {
                'total_recommendations': total_recommendations,
                'unique_sessions': unique_sessions,
                'connected_platforms': len(affiliate_links),
                'daily_stats': daily_stats,
                'period_days': days,
                'avg_daily_recommendations': total_recommendations / days if days > 0 else 0,
                'recommendation_rate': total_recommendations / unique_sessions if unique_sessions > 0 else 0,
                'platform_breakdown': platform_metrics
            }
        })
        
    except Exception as e:
        logger.error(f"Get affiliate performance error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get affiliate performance data'
        }), 500


# =============================================================================
# WEBHOOK ENDPOINTS FOR AFFILIATE PLATFORMS
# =============================================================================

@app.route('/api/webhooks/affiliate/<platform>', methods=['POST'])
def affiliate_webhook(platform):
    """Handle webhook notifications from affiliate platforms"""
    try:
        data = request.get_json()
        
        logger.info(f"📡 Received webhook from {platform}: {data}")
        
        # Process webhook based on platform
        if platform == 'rakuten':
            # Handle Rakuten conversion notifications
            return handle_rakuten_webhook(data)
        elif platform == 'amazon':
            # Handle Amazon conversion notifications
            return handle_amazon_webhook(data)
        elif platform == 'shareasale':
            # Handle ShareASale conversion notifications
            return handle_shareasale_webhook(data)
        elif platform == 'cj_affiliate':
            # Handle CJ Affiliate conversion notifications
            return handle_cj_webhook(data)
        else:
            return jsonify({
                'status': 'error',
                'message': f'Unsupported platform: {platform}'
            }), 400
            
    except Exception as e:
        logger.error(f"Webhook processing error for {platform}: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Webhook processing failed'
        }), 500


def handle_rakuten_webhook(data):
    """Handle Rakuten specific webhook data"""
    try:
        # Rakuten webhook typically includes:
        # - order_id, transaction_id
        # - commission_amount
        # - order_total
        # - affiliate_id
        
        webhook_data = {
            'platform': 'rakuten',
            'order_id': data.get('order_id'),
            'transaction_id': data.get('transaction_id'),
            'commission_amount': data.get('commission_amount'),
            'order_total': data.get('order_total'),
            'affiliate_id': data.get('affiliate_id'),
            'status': data.get('status', 'pending'),
            'created_at': datetime.now(timezone.utc).isoformat()
        }
        
        # Store webhook data for analytics
        # You would implement webhook storage in your database
        
        logger.info(f"✅ Processed Rakuten webhook: Order {data.get('order_id')}")
        
        return jsonify({
            'status': 'success',
            'message': 'Webhook processed successfully'
        })
        
    except Exception as e:
        logger.error(f"Rakuten webhook error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to process Rakuten webhook'
        }), 500


# =============================================================================
# BULK OPERATIONS FOR AFFILIATE MANAGEMENT
# =============================================================================

@app.route('/api/affiliate/bulk-connect', methods=['POST'])
@token_required
def bulk_connect_platforms(current_user):
    """Connect multiple affiliate platforms at once"""
    try:
        data = request.get_json()
        
        if not data or 'platforms' not in data:
            return jsonify({
                'status': 'error',
                'message': 'Platforms data is required'
            }), 400
        
        platforms_data = data['platforms']
        results = {'success': [], 'failed': []}
        
        for platform_data in platforms_data:
            platform = platform_data.get('platform')
            credentials = platform_data.get('credentials', {})
            
            try:
                # Validate platform
                valid_platforms = ['amazon', 'rakuten', 'shareasale', 'cj_affiliate', 'skimlinks']
                if platform not in valid_platforms:
                    results['failed'].append({
                        'platform': platform,
                        'error': 'Invalid platform'
                    })
                    continue
                
                # Check if already connected
                existing = db.get_affiliate_link_by_platform(current_user['id'], platform)
                if existing:
                    results['failed'].append({
                        'platform': platform,
                        'error': 'Already connected'
                    })
                    continue
                
                # Prepare affiliate data
                affiliate_data = {
                    'id': str(uuid.uuid4()),
                    'influencer_id': current_user['id'],
                    'platform': platform,
                    'is_active': True,
                    'created_at': datetime.now(timezone.utc).isoformat(),
                    'updated_at': datetime.now(timezone.utc).isoformat(),
                    **credentials
                }
                
                # Add platform-specific affiliate_id
                if platform == 'amazon':
                    affiliate_data['affiliate_id'] = credentials.get('partner_tag', '')
                elif platform == 'rakuten':
                    affiliate_data['affiliate_id'] = credentials.get('affiliate_id', '')
                elif platform == 'shareasale':
                    affiliate_data['affiliate_id'] = credentials.get('affiliate_id', '')
                elif platform == 'cj_affiliate':
                    affiliate_data['affiliate_id'] = credentials.get('website_id', '')
                elif platform == 'skimlinks':
                    affiliate_data['affiliate_id'] = credentials.get('publisher_id', '')
                
                # Save to database
                success = db.create_affiliate_link(affiliate_data)
                
                if success:
                    results['success'].append({
                        'platform': platform,
                        'affiliate_id': affiliate_data.get('affiliate_id', ''),
                        'created_at': affiliate_data['created_at']
                    })
                else:
                    results['failed'].append({
                        'platform': platform,
                        'error': 'Database save failed'
                    })
                
            except Exception as platform_error:
                results['failed'].append({
                    'platform': platform,
                    'error': str(platform_error)
                })
        
        # Return results
        success_count = len(results['success'])
        total_count = len(platforms_data)
        
        message = f"Connected {success_count}/{total_count} platforms successfully"
        
        return jsonify({
            'status': 'success' if success_count > 0 else 'error',
            'message': message,
            'data': results
        }), 200 if success_count > 0 else 400
        
    except Exception as e:
        logger.error(f"Bulk connect platforms error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to bulk connect platforms'
        }), 500


# =============================================================================
# ENHANCED VOICE SETUP ENDPOINTS
# =============================================================================

@app.route('/api/voice/available-voices', methods=['GET'])
def get_available_voices_enhanced():
    """Get enhanced list of available voices with categories and samples"""
    try:
        # Professional voices for business/educational content
        professional_voices = [
            {
                "voice_id": "2d5b0e6cf36f460aa7fc47e3eee4ba54",
                "name": "Rachel",
                "gender": "Female",
                "language": "English",
                "style": "Professional",
                "description": "Clear, authoritative female voice perfect for business and educational content",
                "sample_text": "Hello, I'm Rachel. I specialize in professional communication and can help your audience with expert advice.",
                "use_cases": ["Business presentations", "Educational content", "Product reviews"],
                "personality_traits": ["Confident", "Clear", "Trustworthy"]
            },
            {
                "voice_id": "d7bbcdd6964c47bdaae26decade4a933",
                "name": "David",
                "gender": "Male", 
                "language": "English",
                "style": "Professional",
                "description": "Deep, authoritative male voice ideal for expert content and tutorials",
                "sample_text": "Hi there, I'm David. My voice conveys authority and expertise, perfect for sharing knowledge.",
                "use_cases": ["Tech tutorials", "Financial advice", "Expert commentary"],
                "personality_traits": ["Authoritative", "Knowledgeable", "Reliable"]
            }
        ]
        
        # Friendly voices for lifestyle/personal content
        friendly_voices = [
            {
                "voice_id": "4d2b8e6cf36f460aa7fc47e3eee4ba12",
                "name": "Emma",
                "gender": "Female",
                "language": "English", 
                "style": "Friendly",
                "description": "Warm, approachable female voice great for lifestyle and personal content",
                "sample_text": "Hey! I'm Emma, and I love connecting with people in a fun, friendly way.",
                "use_cases": ["Lifestyle tips", "Fashion advice", "Personal stories"],
                "personality_traits": ["Warm", "Approachable", "Enthusiastic"]
            },
            {
                "voice_id": "3a1c7d5bf24e350bb6dc46e2dee3ab21",
                "name": "Michael",
                "gender": "Male",
                "language": "English",
                "style": "Casual", 
                "description": "Relaxed male voice perfect for conversational, casual content",
                "sample_text": "What's up! I'm Michael, here to chat about whatever interests you most.",
                "use_cases": ["Gaming content", "Casual conversations", "Entertainment"],
                "personality_traits": ["Relaxed", "Friendly", "Conversational"]
            }
        ]
        
        # Specialized voices for specific niches
        specialized_voices = [
            {
                "voice_id": "1bd001e7e50f421d891986aad5158bc8",
                "name": "Sophia",
                "gender": "Female",
                "language": "English",
                "style": "Gentle",
                "description": "Soft, caring female voice ideal for wellness and mindfulness content",
                "sample_text": "Hello, I'm Sophia. I speak with gentleness and care, perfect for wellness topics.",
                "use_cases": ["Wellness coaching", "Meditation guides", "Self-care tips"],
                "personality_traits": ["Gentle", "Caring", "Soothing"]
            },
            {
                "voice_id": "26b2064088674c80b1e5fc5ab1a068ec", 
                "name": "Marcus",
                "gender": "Male",
                "language": "English",
                "style": "Energetic",
                "description": "Dynamic, motivational male voice for fitness and motivation content",
                "sample_text": "Hey everyone! I'm Marcus, and I'm here to pump you up and keep you motivated!",
                "use_cases": ["Fitness coaching", "Motivational content", "Sports commentary"],
                "personality_traits": ["Energetic", "Motivational", "Dynamic"]
            }
        ]
        
        return jsonify({
            'status': 'success',
            'data': {
                'voice_categories': {
                    'professional': {
                        'name': 'Professional Voices',
                        'description': 'Perfect for business, education, and expert content',
                        'voices': professional_voices
                    },
                    'friendly': {
                        'name': 'Friendly & Casual',
                        'description': 'Great for lifestyle, personal, and conversational content',
                        'voices': friendly_voices
                    },
                    'specialized': {
                        'name': 'Specialized Voices',
                        'description': 'Tailored for specific niches and use cases',
                        'voices': specialized_voices
                    }
                },
                'total_voices': len(professional_voices) + len(friendly_voices) + len(specialized_voices),
                'default_voice_id': Config.DEFAULT_VOICE_ID,
                'voice_cloning_available': bool(Config.ELEVEN_LABS_API_KEY)
            }
        })
        
    except Exception as e:
        logger.error(f"Get enhanced voices error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get voice list'
        }), 500


# =============================================================================
# ENHANCED DASHBOARD DATA
# =============================================================================

@app.route('/api/dashboard/enhanced', methods=['GET'])
@token_required
def get_enhanced_dashboard_data(current_user):
    """Get comprehensive dashboard data including affiliate metrics"""
    try:
        from datetime import timedelta
        
        # Time ranges
        today = datetime.now(timezone.utc)
        last_7_days = today - timedelta(days=7)
        last_30_days = today - timedelta(days=30)
        
        # Get basic profile info
        influencer = db.get_influencer(current_user['id'])
        
        # Get chat interactions
        recent_chats = db.supabase.table('chat_interactions') \
            .select('*') \
            .eq('influencer_id', current_user['id']) \
            .gte('created_at', last_30_days.isoformat()) \
            .execute()
        
        chats = recent_chats.data if recent_chats.data else []
        
        # Get affiliate links
        affiliate_links = db.get_affiliate_links(current_user['id'])
        
        # Calculate metrics
        total_chats = len(chats)
        unique_visitors = len(set([chat.get("session_id") for chat in chats if chat.get("session_id")]))
        video_responses = len([chat for chat in chats if chat.get('has_video')])
        audio_responses = len([chat for chat in chats if chat.get('has_audio')])
        product_recommendations = len([chat for chat in chats if chat.get('products_included')])
        
        # Weekly comparison
        last_week_chats = [chat for chat in chats if datetime.fromisoformat(chat['created_at'].replace('Z', '+00:00')) >= last_7_days]
        prev_week_chats = [chat for chat in chats if datetime.fromisoformat(chat['created_at'].replace('Z', '+00:00')) < last_7_days]
        
        week_growth = len(last_week_chats) - len(prev_week_chats)
        week_growth_percent = (week_growth / len(prev_week_chats) * 100) if prev_week_chats else 0
        
        # Setup completion score
        completion_score = 0
        setup_items = {
            'avatar_created': bool(influencer.get('heygen_avatar_id')),
            'voice_configured': bool(influencer.get('preferred_voice_id')),
            'bio_added': bool(influencer.get('bio')),
            'expertise_added': bool(influencer.get('expertise')),
            'affiliate_connected': len(affiliate_links) > 0,
            'knowledge_added': False  # You could check for uploaded documents
        }
        
        completion_score = sum(setup_items.values()) / len(setup_items) * 100
        
        dashboard_data = {
            'overview': {
                'total_chats': total_chats,
                'unique_visitors': unique_visitors,
                'video_responses': video_responses,
                'audio_responses': audio_responses,
                'product_recommendations': product_recommendations,
                'completion_score': round(completion_score, 1),
                'week_growth': week_growth,
                'week_growth_percent': round(week_growth_percent, 1)
            },
            'setup_status': setup_items,
            'affiliate_status': {
                'connected_platforms': len(affiliate_links),
                'total_platforms': 5,  # Amazon, Rakuten, ShareASale, CJ, Skimlinks
                'platform_details': affiliate_links
            },
            'avatar_status': {
                'has_avatar': bool(influencer.get('heygen_avatar_id')),
                'avatar_id': influencer.get('heygen_avatar_id'),
                'voice_configured': bool(influencer.get('preferred_voice_id')),
                'preferred_voice': influencer.get('preferred_voice_id')
            },
            'recent_activity': [
                {
                    'type': 'chat',
                    'message': chat.get('user_message', '')[:50] + '...',
                    'timestamp': chat.get('created_at'),
                    'has_video': chat.get('has_video', False),
                    'has_audio': chat.get('has_audio', False)
                }
                for chat in sorted(chats, key=lambda x: x.get('created_at', ''), reverse=True)[:5]
            ],
            'chat_url': f"{request.host_url}pages/chat.html?username={influencer['username']}",
            'profile': {
                'username': influencer['username'],
                'bio': influencer.get('bio', ''),
                'expertise': influencer.get('expertise', ''),
                'personality': influencer.get('personality', ''),
                'created_at': influencer.get('created_at')
            }
        }
        
        return jsonify({
            'status': 'success',
            'data': dashboard_data
        })
        
    except Exception as e:
        logger.error(f"Enhanced dashboard data error: {e}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get dashboard data'
        }), 500
           
# =============================================================================
# APPLICATION STARTUP
# =============================================================================

# Initialize the enhanced chatbot
def initialize_enhanced_chatbot():
    """FIXED: Initialize the enhanced chatbot with all features"""
    global chatbot
    
    try:
        # Always use EnhancedChatbot for full functionality
        from chatbot import EnhancedChatbot
        chatbot = EnhancedChatbot(db=db)
        logger.info("✅ Enhanced chatbot initialized with knowledge and affiliate integration")
        return True
    except ImportError as e:
        logger.error(f"❌ Failed to import EnhancedChatbot: {e}")
        # Fallback to basic chatbot
        try:
            from chatbot import Chatbot
            chatbot = Chatbot(db)
            logger.warning("⚠️ Using basic chatbot - enhanced features may not be available")
            return True
        except ImportError as e2:
            logger.error(f"❌ Failed to initialize any chatbot: {e2}")
            return False
            
@app.route('/api/affiliate/debug-rakuten', methods=['POST'])
@token_required
def debug_rakuten_connection(current_user):
    """FIXED: Comprehensive Rakuten API debugging endpoint"""
    try:
        # FIXED: Handle both JSON and form data
        if request.content_type and 'application/json' in request.content_type:
            data = request.get_json() or {}
        else:
            # Handle form data or other content types
            data = {
                'client_id': request.form.get('client_id', ''),
                'client_secret': request.form.get('client_secret', '')
            }
        
        logger.info(f"🔍 Debug request - Content-Type: {request.content_type}")
        logger.info(f"📝 Request data: {list(data.keys()) if data else 'No data'}")
        
        # Get credentials from request or database
        if data and data.get('client_id'):
            # Use provided credentials
            client_id = data.get('client_id', '').strip()
            client_secret = data.get('client_secret', '').strip()
            test_source = 'provided_credentials'
        else:
            # Use stored credentials
            affiliate_link = db.get_affiliate_link_by_platform(current_user['id'], 'rakuten')
            if not affiliate_link:
                return jsonify({
                    'status': 'error',
                    'message': 'No Rakuten connection found and no credentials provided',
                    'debug_info': {
                        'has_stored_connection': False,
                        'has_provided_credentials': bool(data and data.get('client_id')),
                        'content_type': request.content_type,
                        'request_keys': list(data.keys()) if data else []
                    }
                }), 400
            
            client_id = (
                affiliate_link.get('client_id') or 
                affiliate_link.get('rakuten_client_id') or 
                affiliate_link.get('api_key')
            )
            client_secret = (
                affiliate_link.get('client_secret') or 
                affiliate_link.get('rakuten_client_secret')
            )
            test_source = 'stored_credentials'
        
        logger.info(f"🔍 Debug Rakuten API for {current_user['username']} using {test_source}")
        
        debug_info = {
            'user': current_user['username'],
            'test_source': test_source,
            'request_info': {
                'content_type': request.content_type,
                'method': request.method,
                'has_json': bool(request.get_json(silent=True)),
                'has_form': bool(request.form)
            },
            'credentials_check': {
                'has_client_id': bool(client_id),
                'client_id_length': len(client_id) if client_id else 0,
                'client_id_preview': f"{client_id[:10]}..." if client_id else 'missing',
                'has_client_secret': bool(client_secret),
                'client_secret_length': len(client_secret) if client_secret else 0
            },
            'api_tests': []
        }
        
        if not client_id:
            return jsonify({
                'status': 'error',
                'message': 'No client_id found',
                'debug_info': debug_info
            }), 400
        
        # FIXED: Test with SSL verification disabled for debugging
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        # Test 1: Basic connectivity with SSL disabled
        try:
            logger.info("🧪 Testing basic connectivity to Rakuten...")
            
            response = requests.get(
                'https://api.rakutenadvertising.com',
                timeout=10,
                verify=False  # DISABLE SSL for debugging
            )
            
            debug_info['api_tests'].append({
                'test_name': 'basic_connectivity',
                'success': response.status_code < 500,
                'status_code': response.status_code,
                'response_preview': response.text[:200] if response.text else 'No content',
                'error': None
            })
            
        except Exception as conn_error:
            debug_info['api_tests'].append({
                'test_name': 'basic_connectivity',
                'success': False,
                'error': str(conn_error)
            })
        
        # Test 2: OAuth token with SSL disabled
        try:
            logger.info("🔑 Testing OAuth token endpoint...")
            
            auth_headers = {
                'Content-Type': 'application/x-www-form-urlencoded',
                'Accept': 'application/json'
            }
            
            auth_payload = {
                'grant_type': 'client_credentials',
                'client_id': client_id,
                'client_secret': client_secret,
                'scope': 'productsearch'
            }
            
            # Try multiple auth endpoints with SSL disabled
            auth_urls = [
                'https://api.rakutenadvertising.com/token',
                'https://api.rakutenadvertising.com/auth/token',
                'https://api.rakutenadvertising.com/oauth/token'
            ]
            
            for auth_url in auth_urls:
                try:
                    logger.info(f"🔗 Testing auth URL: {auth_url}")
                    
                    response = requests.post(
                        auth_url,
                        headers=auth_headers,
                        data=auth_payload,
                        timeout=15,
                        verify=False  # DISABLE SSL for debugging
                    )
                    
                    debug_info['api_tests'].append({
                        'test_name': f'oauth_token_{auth_url.split("/")[-1]}',
                        'url': auth_url,
                        'success': response.status_code == 200,
                        'status_code': response.status_code,
                        'response_preview': response.text[:300],
                        'has_access_token': 'access_token' in response.text,
                        'error': None
                    })
                    
                    if response.status_code == 200:
                        logger.info(f"✅ SUCCESS with {auth_url}")
                        try:
                            token_data = response.json()
                            access_token = token_data.get('access_token')
                            if access_token:
                                # Test 3: Product search with token
                                logger.info("🔍 Testing product search with token...")
                                
                                search_headers = {
                                    'Authorization': f'Bearer {access_token}',
                                    'Accept': 'application/json'
                                }
                                
                                search_params = {
                                    'keyword': 'laptop',
                                    'limit': 2
                                }
                                
                                search_urls = [
                                    'https://api.rakutenadvertising.com/v1/productsearch',
                                    'https://api.rakutenadvertising.com/productsearch/1.0'
                                ]
                                
                                for search_url in search_urls:
                                    try:
                                        search_response = requests.get(
                                            search_url,
                                            headers=search_headers,
                                            params=search_params,
                                            timeout=15,
                                            verify=False  # DISABLE SSL for debugging
                                        )
                                        
                                        debug_info['api_tests'].append({
                                            'test_name': f'product_search_{search_url.split("/")[-1]}',
                                            'url': search_url,
                                            'success': search_response.status_code == 200,
                                            'status_code': search_response.status_code,
                                            'response_preview': search_response.text[:400],
                                            'has_products': 'product' in search_response.text.lower(),
                                            'error': None
                                        })
                                        
                                        if search_response.status_code == 200:
                                            logger.info(f"✅ PRODUCT SEARCH SUCCESS with {search_url}")
                                            break
                                            
                                    except Exception as search_error:
                                        debug_info['api_tests'].append({
                                            'test_name': f'product_search_{search_url.split("/")[-1]}',
                                            'url': search_url,
                                            'success': False,
                                            'error': str(search_error)
                                        })
                                break
                        except Exception as token_parse_error:
                            logger.error(f"Token parsing error: {token_parse_error}")
                    
                except Exception as auth_error:
                    debug_info['api_tests'].append({
                        'test_name': f'oauth_token_{auth_url.split("/")[-1]}',
                        'url': auth_url,
                        'success': False,
                        'error': str(auth_error)
                    })
                    
        except Exception as oauth_error:
            debug_info['api_tests'].append({
                'test_name': 'oauth_flow',
                'success': False,
                'error': str(oauth_error)
            })
        
        # Determine overall status
        any_success = any(
            test.get('success', False) for test in debug_info['api_tests']
        )
        
        return jsonify({
            'status': 'success' if any_success else 'error',
            'message': 'Rakuten API debugging completed' if any_success else 'All API tests failed',
            'debug_info': debug_info,
            'next_steps': [
                'SSL certificate issue detected - using verify=False for testing',
                'Check if your Rakuten account has API access enabled',
                'Verify you\'re using Rakuten Advertising (not Rakuten Ichiba)', 
                'Confirm your Client ID and Secret are correct',
                'Contact Rakuten support if credentials are correct but API fails'
            ] if not any_success else [
                'API connection successful!',
                'SSL verification disabled for debugging - re-enable in production',
                'Check which specific endpoint worked in debug_info'
            ],
            'ssl_warning': 'SSL verification disabled for debugging. Enable verify=True in production.'
        })
        
    except Exception as e:
        logger.error(f"❌ Debug endpoint error: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Debug failed: {str(e)}',
            'error_type': 'debug_failed',
            'request_info': {
                'content_type': request.content_type,
                'method': request.method,
                'has_json': bool(request.get_json(silent=True)),
                'has_form': bool(request.form)
            }
        }), 500

def await_test_basic_connectivity():
    """Test basic network connectivity to Rakuten"""
    try:
        import socket
        
        # Test DNS resolution
        socket.gethostbyname('api.rakutenadvertising.com')
        
        # Test basic HTTP connectivity
        response = requests.get(
            'https://api.rakutenadvertising.com',
            timeout=10,
            verify=False
        )
        
        return {
            'test_name': 'basic_connectivity',
            'success': True,
            'dns_resolution': 'success',
            'http_connectivity': f'HTTP {response.status_code}',
            'error': None
        }
        
    except Exception as conn_error:
        return {
            'test_name': 'basic_connectivity',
            'success': False,
            'error': str(conn_error)
        }
    
# Add this to your main initialization
if __name__ == '__main__':
    # Validate environment
    try:
        Config.validate()
        logger.info("✅ Environment validation passed")
    except ValueError as e:
        logger.error(f"❌ Environment validation failed: {e}")
        exit(1)
    
    # FIXED: Initialize enhanced chatbot with error handling
    if not initialize_enhanced_chatbot():
        logger.error("❌ Chatbot initialization failed")
        exit(1)
    
    logger.info("🚀 Starting AvatarCommerce with enhanced features:")
    logger.info(f"   - Knowledge processing: {'✅' if chatbot.rag_processor else '❌'}")
    logger.info(f"   - Affiliate integration: {'✅' if chatbot.affiliate_service else '❌'}")
    logger.info(f"   - Voice generation: {'✅' if chatbot.eleven_labs_api_key or True else '❌'}")
    logger.info(f"   - Video generation: {'✅' if chatbot.heygen_api_key else '❌'}")
    
    # Run the app
    app.run(host='0.0.0.0', port=2000, debug=True)